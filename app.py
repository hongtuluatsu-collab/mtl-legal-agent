"""
=============================================================
  MTL LEGAL AGENT PREMIUM
  Công ty Luật TNHH Minh Tú
=============================================================
Cài đặt (chạy 1 lần):
  pip install streamlit anthropic python-docx PyPDF2 pillow

Chạy app:
  streamlit run app.py
=============================================================
"""

import streamlit as st
import anthropic
import base64
import io
import os
import re
import json
from datetime import datetime

# OAuth + Gmail API
from google_auth_oauthlib.flow import Flow
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
import urllib.parse
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
from calendar import monthrange

# =============================================================
#  API KEY — cấu hình trong Railway Variables
#  Tên biến trong Railway: ANTHROPIC_API_KEY
# =============================================================
API_KEY_FALLBACK = ""
# =============================================================

# ─────────────────────────────────────────────
#  THÔNG TIN CÔNG TY
# ─────────────────────────────────────────────
TEN_CONG_TY = "CÔNG TY LUẬT TNHH MINH TÚ"
DIA_CHI_CT  = "Trụ sở: 4/9 Đường số 3 Cư Xá Đô Thành, Phường Bàn Cờ, TP. Hồ Chí Minh"
DIA_CHI_DN  = "CN Đà Nẵng: 81 Xô Viết Nghệ Tĩnh, Phường Cẩm Lệ, TP. Đà Nẵng"
SBT_CT      = "Hotline: 19000031 | Website: luatminhtu.vn | Email: info.luatminhtu@gmail.com"

# ─────────────────────────────────────────────
#  TÀI KHOẢN LUẬT SƯ
# ─────────────────────────────────────────────
TAI_KHOAN = {
    "ls.hoang": {
        "mat_khau": "Hoang@2026",
        "ho_ten":   "Luật sư Nguyễn Minh Hoàng",
        "chuc_vu":  "Luật sư thành viên",
        "vai_tro":  "luat_su",
    },
    "ls.thang": {
        "mat_khau": "Thang@2026",
        "ho_ten":   "Luật sư Trịnh Chiến Thắng",
        "chuc_vu":  "Luật sư thành viên",
        "vai_tro":  "luat_su",
    },
    "ls.lan": {
        "mat_khau": "Lan@2026",
        "ho_ten":   "Luật sư Nguyễn Thị Thanh Lan",
        "chuc_vu":  "Luật sư thành viên",
        "vai_tro":  "luat_su",
    },
    "ls.phuong": {
        "mat_khau": "Phuong@2026",
        "ho_ten":   "Luật sư Lê Thuý Phượng",
        "chuc_vu":  "Luật sư thành viên",
        "vai_tro":  "luat_su",
    },
    "ls.nga": {
        "mat_khau": "Nga@2026",
        "ho_ten":   "Luật sư Phạm Thị Nga",
        "chuc_vu":  "Luật sư thành viên",
        "vai_tro":  "luat_su",
    },
    "ls.dong": {
        "mat_khau": "Dong@2026",
        "ho_ten":   "Luật sư Lê Viễn Đông",
        "chuc_vu":  "Luật sư thành viên",
        "vai_tro":  "luat_su",
    },
    "admin": {
        "mat_khau": "Admin@MTL2026",
        "ho_ten":   "Quản trị viên",
        "chuc_vu":  "Quản lý hệ thống",
        "vai_tro":  "quan_tri",
    },
}

# ─────────────────────────────────────────────
#  MÀU THƯƠNG HIỆU MTL
# ─────────────────────────────────────────────
MTL_NAVY  = "#1E4D82"
MTL_GOLD  = "#A8874A"
MTL_NAVY2 = "#163960"
MTL_GOLD2 = "#C9A96E"

# ─────────────────────────────────────────────
#  CẤU HÌNH TRANG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="MTL Legal Agent Premium",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(f"""
<style>
#MainMenu, footer, header {{ visibility: hidden; }}

section[data-testid="stSidebar"] {{
    background: linear-gradient(180deg, {MTL_NAVY2} 0%, {MTL_NAVY} 60%, #122d50 100%);
    border-right: 2px solid {MTL_GOLD};
}}
section[data-testid="stSidebar"] * {{ color: #e8eef5 !important; }}
section[data-testid="stSidebar"] input {{
    background: rgba(255,255,255,0.08) !important;
    border: 1px solid {MTL_GOLD} !important;
    color: white !important;
    border-radius: 6px !important;
}}
section[data-testid="stSidebar"] hr {{ border-color: {MTL_GOLD}44 !important; }}
section[data-testid="stSidebar"] .stButton > button {{
    background: transparent !important;
    border: 1px solid {MTL_GOLD} !important;
    color: {MTL_GOLD} !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
}}
section[data-testid="stSidebar"] .stButton > button:hover {{
    background: {MTL_GOLD} !important;
    color: white !important;
}}
section[data-testid="stSidebar"] .stFileUploader [data-testid="stFileUploaderDropzone"] {{
    background: rgba(255,255,255,0.05) !important;
    border: 2px dashed {MTL_GOLD}99 !important;
    border-radius: 10px !important;
}}
section[data-testid="stSidebar"] .stFileUploader [data-testid="stFileUploaderDropzone"] * {{
    color: {MTL_GOLD2} !important;
}}
section[data-testid="stSidebar"] .stFileUploader button {{
    background: {MTL_GOLD} !important;
    color: white !important;
    border: none !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
}}
.mtl-header {{
    background: linear-gradient(135deg, {MTL_NAVY2} 0%, {MTL_NAVY} 70%, #1a5592 100%);
    border-bottom: 3px solid {MTL_GOLD};
    border-radius: 12px;
    margin-bottom: 24px;
    overflow: hidden;
    box-shadow: 0 4px 20px rgba(30,77,130,0.25);
}}
.mtl-header-inner {{
    display: flex; align-items: center; gap: 20px; padding: 18px 28px;
}}
.mtl-box {{
    width: 40px; height: 40px; border-radius: 6px;
    display: flex; align-items: center; justify-content: center;
    font-size: 1.2rem; font-weight: 700; color: white; font-family: Georgia, serif;
}}
.mtl-box-navy {{ background: {MTL_NAVY2}; border: 1.5px solid rgba(255,255,255,0.3); }}
.mtl-box-gold  {{ background: {MTL_GOLD};  border: 1.5px solid rgba(255,255,255,0.3); }}
.mtl-divider {{
    width: 2px; height: 52px;
    background: linear-gradient(to bottom, transparent, {MTL_GOLD}, transparent);
    margin: 0 12px;
}}
.mtl-title-block h1 {{ margin: 0; color: white; font-size: 1.25rem; font-weight: 700; }}
.mtl-title-block .mtl-sub {{ margin: 3px 0 0; color: {MTL_GOLD2}; font-size: 0.82rem; }}
.mtl-user-badge {{
    margin-left: auto;
    background: rgba(255,255,255,0.08);
    border: 1px solid {MTL_GOLD}66;
    border-radius: 8px; padding: 8px 16px; text-align: right;
}}
.mtl-user-badge .name {{ color: white; font-weight: 600; font-size: 0.9rem; }}
.mtl-user-badge .role {{ color: {MTL_GOLD2}; font-size: 0.78rem; }}
.mtl-user-badge .date {{ color: rgba(255,255,255,0.5); font-size: 0.72rem; margin-top: 2px; }}
.login-card {{
    background: white; border-radius: 16px; padding: 40px 36px;
    box-shadow: 0 8px 40px rgba(30,77,130,0.15);
    border-top: 4px solid {MTL_GOLD}; max-width: 420px; margin: 0 auto;
}}
.login-title {{ text-align: center; color: {MTL_NAVY}; font-size: 1.3rem; font-weight: 700; margin: 8px 0 4px; }}
.login-sub {{ text-align: center; color: {MTL_GOLD}; font-size: 0.82rem; letter-spacing: 1px; margin: 0 0 24px; text-transform: uppercase; font-weight: 600; }}
.result-box {{
    background: linear-gradient(135deg, #f0f5ff 0%, #fafbff 100%);
    border-left: 4px solid {MTL_NAVY};
    border-top: 1px solid #e0e8f5; border-right: 1px solid #e0e8f5; border-bottom: 1px solid #e0e8f5;
    padding: 20px 24px; border-radius: 0 10px 10px 0; margin-top: 16px; line-height: 1.75;
}}
.stTabs [data-baseweb="tab-list"] {{ gap: 4px; border-bottom: 2px solid {MTL_GOLD}44; }}
.stTabs [data-baseweb="tab"] {{ border-radius: 8px 8px 0 0 !important; font-weight: 600 !important; padding: 8px 18px !important; }}
.stTabs [aria-selected="true"] {{ background: {MTL_NAVY} !important; color: white !important; border-bottom: 2px solid {MTL_GOLD} !important; }}
.stButton > button {{ border-radius: 8px !important; font-weight: 600 !important; transition: all 0.2s !important; }}
.stButton > button[kind="primary"] {{
    background: linear-gradient(135deg, {MTL_NAVY} 0%, {MTL_NAVY2} 100%) !important;
    border: none !important; color: white !important;
    box-shadow: 0 2px 8px rgba(30,77,130,0.3) !important;
}}
.stDownloadButton > button {{
    background: linear-gradient(135deg, {MTL_GOLD} 0%, #8a6d38 100%) !important;
    color: white !important; border: none !important;
    border-radius: 8px !important; font-weight: 600 !important;
}}

/* ── Sidebar toggle button ── */
#mtl-sidebar-toggle {{
    position: fixed;
    top: 50%;
    left: 0;
    transform: translateY(-50%);
    z-index: 9999;
    width: 18px;
    height: 64px;
    background: {MTL_NAVY};
    border: 1.5px solid {MTL_GOLD};
    border-left: none;
    border-radius: 0 8px 8px 0;
    cursor: pointer;
    display: flex;
    align-items: center;
    justify-content: center;
    box-shadow: 2px 0 10px rgba(30,77,130,0.25);
    transition: width 0.2s, background 0.2s;
}}
#mtl-sidebar-toggle:hover {{
    width: 24px;
    background: {MTL_GOLD};
}}
#mtl-sidebar-toggle .mtl-arrow {{
    color: {MTL_GOLD};
    font-size: 11px;
    font-weight: 700;
    line-height: 1;
    user-select: none;
    transition: color 0.2s;
    pointer-events: none;
}}
#mtl-sidebar-toggle:hover .mtl-arrow {{
    color: white;
}}

/* Trạng thái sidebar thu nhỏ */
body.sidebar-collapsed section[data-testid="stSidebar"] {{
    width: 0 !important;
    min-width: 0 !important;
    overflow: hidden !important;
    visibility: hidden;
    opacity: 0;
    transition: all 0.3s ease;
}}
body.sidebar-collapsed section[data-testid="stSidebar"] + div,
body.sidebar-collapsed .main {{
    margin-left: 0 !important;
    padding-left: 0 !important;
}}
section[data-testid="stSidebar"] {{
    transition: all 0.3s ease;
}}
</style>
""", unsafe_allow_html=True)

# ── Inject toggle button + JS (chạy 1 lần sau khi CSS load) ──
st.markdown(f"""
<div id="mtl-sidebar-toggle" onclick="mtlToggleSidebar()" title="Thu/mở thanh bên">
  <span class="mtl-arrow" id="mtl-arrow-icon">&#8249;</span>
</div>

<script>
(function() {{
  // Khôi phục trạng thái từ localStorage
  var collapsed = localStorage.getItem('mtl_sidebar_collapsed') === '1';
  if (collapsed) {{
    document.body.classList.add('sidebar-collapsed');
    var icon = document.getElementById('mtl-arrow-icon');
    if (icon) icon.innerHTML = '&#8250;';
    // Dịch nút sang trái khi sidebar thu
    var btn = document.getElementById('mtl-sidebar-toggle');
    if (btn) btn.style.left = '0';
  }}
}})();

window.mtlToggleSidebar = function() {{
  var body     = document.body;
  var btn      = document.getElementById('mtl-sidebar-toggle');
  var icon     = document.getElementById('mtl-arrow-icon');
  var sidebar  = document.querySelector('section[data-testid="stSidebar"]');
  var collapsed = body.classList.toggle('sidebar-collapsed');

  if (collapsed) {{
    // Thu nhỏ
    icon.innerHTML = '&#8250;';
    btn.style.left = '0px';
    if (sidebar) {{
      sidebar.style.width = '0';
      sidebar.style.minWidth = '0';
      sidebar.style.overflow = 'hidden';
      sidebar.style.visibility = 'hidden';
      sidebar.style.opacity = '0';
    }}
    localStorage.setItem('mtl_sidebar_collapsed', '1');
  }} else {{
    // Mở rộng
    icon.innerHTML = '&#8249;';
    if (sidebar) {{
      sidebar.style.width = '';
      sidebar.style.minWidth = '';
      sidebar.style.overflow = '';
      sidebar.style.visibility = '';
      sidebar.style.opacity = '';
      // Lấy lại vị trí nút sau khi sidebar hiện
      setTimeout(function() {{
        var sw = sidebar.offsetWidth || 300;
        btn.style.left = sw + 'px';
      }}, 320);
    }}
    localStorage.setItem('mtl_sidebar_collapsed', '0');
  }}
}};

// Gắn vị trí nút theo chiều rộng sidebar thực tế
(function positionBtn() {{
  var sidebar = document.querySelector('section[data-testid="stSidebar"]');
  var btn     = document.getElementById('mtl-sidebar-toggle');
  if (!sidebar || !btn) {{ setTimeout(positionBtn, 200); return; }}

  var collapsed = localStorage.getItem('mtl_sidebar_collapsed') === '1';
  if (!collapsed) {{
    var sw = sidebar.offsetWidth || 300;
    btn.style.left = sw + 'px';
  }}

  // Theo dõi resize sidebar (Streamlit thay đổi layout)
  var ro = new ResizeObserver(function(entries) {{
    if (!document.body.classList.contains('sidebar-collapsed')) {{
      var sw = entries[0].contentRect.width;
      if (sw > 0) btn.style.left = sw + 'px';
    }}
  }});
  ro.observe(sidebar);
}})();
</script>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  XÁC ĐỊNH API KEY
# ─────────────────────────────────────────────
def lay_api_key():
    # Đọc thẳng từ env var Railway
    key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if key:
        return key
    # Fallback: key dự phòng điền tay
    if API_KEY_FALLBACK and len(API_KEY_FALLBACK) > 20:
        return API_KEY_FALLBACK
    return ""

API_KEY = lay_api_key()

# ─────────────────────────────────────────────
#  ĐĂNG NHẬP
# ─────────────────────────────────────────────
if "dang_nhap" not in st.session_state:
    st.session_state.dang_nhap = False
    st.session_state.nguoi_dung = None

def dang_nhap(ten_tk, mat_khau):
    if ten_tk in TAI_KHOAN and TAI_KHOAN[ten_tk]["mat_khau"] == mat_khau:
        st.session_state.dang_nhap = True
        st.session_state.nguoi_dung = {**TAI_KHOAN[ten_tk], "ten_tk": ten_tk}
        return True
    return False

def dang_xuat():
    st.session_state.dang_nhap = False
    st.session_state.nguoi_dung = None
    st.rerun()

if not st.session_state.dang_nhap:
    _, col, _ = st.columns([1, 1.6, 1])
    with col:
        st.markdown("<br><br>", unsafe_allow_html=True)
        st.markdown(f"""
<div class="login-card">
  <div style="display:flex;align-items:center;justify-content:center;gap:4px;margin-bottom:8px;">
    <div class="mtl-box mtl-box-navy" style="width:42px;height:42px;">M</div>
    <div class="mtl-box mtl-box-gold"  style="width:42px;height:42px;">T</div>
    <div class="mtl-box mtl-box-navy" style="width:42px;height:42px;">L</div>
  </div>
  <div style="height:3px;background:linear-gradient(to right,#163960,#A8874A,#163960);border-radius:2px;margin:10px 40px 12px;"></div>
  <p class="login-title">LUẬT MINH TÚ</p>
  <p class="login-sub">⚖ Legal Agent Premium</p>
</div>
""", unsafe_allow_html=True)

        with st.form("form_dn"):
            ten_tk   = st.text_input("Tên đăng nhập", placeholder="Ví dụ: ls.hoang")
            mat_khau = st.text_input("Mật khẩu", type="password", placeholder="••••••••")
            nut      = st.form_submit_button("🔐  Đăng nhập", use_container_width=True)

        if nut:
            if dang_nhap(ten_tk.strip(), mat_khau):
                st.success("✅ Đăng nhập thành công!")
                st.rerun()
            else:
                st.error("❌ Sai tên đăng nhập hoặc mật khẩu.")

        st.markdown("<p style='text-align:center;color:#bbb;font-size:0.75rem;margin-top:16px;'>© 2026 Công ty Luật TNHH Minh Tú</p>", unsafe_allow_html=True)
        st.markdown(f"""
<div style="margin-top:18px;padding-top:16px;border-top:1px solid #f0e8d8;">
  <p style="text-align:center;font-size:0.7rem;color:{MTL_GOLD};font-weight:700;
  letter-spacing:2px;text-transform:uppercase;margin-bottom:12px;">Giá trị cốt lõi</p>
  <div style="display:flex;justify-content:center;gap:8px;">
    <div style="flex:1;text-align:center;padding:10px 6px;
    background:linear-gradient(135deg,{MTL_NAVY2},{MTL_NAVY});
    border-radius:10px;border:1px solid {MTL_GOLD}55;">
      <div style="font-size:1.3rem;margin-bottom:4px;">🤝</div>
      <div style="font-size:0.68rem;font-weight:800;color:{MTL_GOLD};
      letter-spacing:1px;text-transform:uppercase;">Cam kết</div>
      <div style="font-size:0.6rem;color:rgba(255,255,255,0.55);margin-top:3px;
      line-height:1.4;">Tận tâm phục vụ<br>đến cùng</div>
    </div>
    <div style="flex:1;text-align:center;padding:10px 6px;
    background:linear-gradient(135deg,{MTL_NAVY2},{MTL_NAVY});
    border-radius:10px;border:1px solid {MTL_GOLD}55;">
      <div style="font-size:1.3rem;margin-bottom:4px;">⚖️</div>
      <div style="font-size:0.68rem;font-weight:800;color:{MTL_GOLD};
      letter-spacing:1px;text-transform:uppercase;">Chính trực</div>
      <div style="font-size:0.6rem;color:rgba(255,255,255,0.55);margin-top:3px;
      line-height:1.4;">Minh bạch &<br>đạo đức nghề nghiệp</div>
    </div>
    <div style="flex:1;text-align:center;padding:10px 6px;
    background:linear-gradient(135deg,{MTL_NAVY2},{MTL_NAVY});
    border-radius:10px;border:1px solid {MTL_GOLD}55;">
      <div style="font-size:1.3rem;margin-bottom:4px;">📚</div>
      <div style="font-size:0.68rem;font-weight:800;color:{MTL_GOLD};
      letter-spacing:1px;text-transform:uppercase;">Học hỏi</div>
      <div style="font-size:0.6rem;color:rgba(255,255,255,0.55);margin-top:3px;
      line-height:1.4;">Không ngừng<br>trau dồi kiến thức</div>
    </div>
  </div>
</div>
""", unsafe_allow_html=True)
    st.stop()

# ─────────────────────────────────────────────
#  HÀM ĐỌC FILE
# ─────────────────────────────────────────────
def doc_pdf(file_bytes):
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        noi_dung = []
        for i, trang in enumerate(reader.pages):
            txt = trang.extract_text()
            if txt:
                noi_dung.append(f"[Trang {i+1}]\n{txt}")
        return "\n\n".join(noi_dung) if noi_dung else "⚠️ Không đọc được nội dung PDF."
    except Exception as e:
        return f"Lỗi đọc PDF: {e}"

def doc_docx(file_bytes):
    try:
        doc = Document(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        return f"Lỗi đọc DOCX: {e}"

# ─────────────────────────────────────────────
#  HÀM GỌI CLAUDE
# ─────────────────────────────────────────────
def goi_claude(messages, system_prompt):
    try:
        client = anthropic.Anthropic(api_key=API_KEY)
        response = client.messages.create(
            model="claude-opus-4-5",
            max_tokens=4096,
            system=system_prompt,
            messages=messages,
        )
        return response.content[0].text
    except anthropic.AuthenticationError:
        return "❌ API Key không hợp lệ. Kiểm tra lại dòng ANTHROPIC_API_KEY ở đầu file app.py."
    except Exception as e:
        return f"❌ Lỗi: {e}"

def phan_tich_ho_so(noi_dung_files, yeu_cau=""):
    system = f"""Bạn là chuyên gia pháp lý Việt Nam tại {TEN_CONG_TY}.
Phân tích hồ sơ bằng tiếng Việt, chuyên nghiệp, có cấu trúc rõ ràng gồm:
1. TÓM TẮT VỤ VIỆC  2. CÁC BÊN LIÊN QUAN  3. VẤN ĐỀ PHÁP LÝ CỐT LÕI
4. CĂN CỨ PHÁP LÝ  5. ĐÁNH GIÁ RỦI RO & CƠ HỘI  6. HƯỚNG XỬ LÝ ĐỀ XUẤT  7. TÀI LIỆU CẦN BỔ SUNG"""
    messages = []
    for item in noi_dung_files:
        if item["loai"] == "anh":
            messages.append({"role": "user", "content": [
                {"type": "image", "source": {"type": "base64", "media_type": item["media_type"], "data": item["du_lieu"]}},
                {"type": "text", "text": f"File: {item['ten']}"},
            ]})
        else:
            messages.append({"role": "user", "content": f"File: {item['ten']}\n\n{item['du_lieu']}"})
    cau_hoi = "Hãy phân tích toàn bộ hồ sơ."
    if yeu_cau:
        cau_hoi += f"\n\nYêu cầu thêm: {yeu_cau}"
    messages.append({"role": "user", "content": cau_hoi})
    return goi_claude(messages, system)

def soan_don_tu(loai_don, noi_dung, them=""):
    system = f"""Bạn là luật sư chuyên nghiệp tại {TEN_CONG_TY}.
Soạn {loai_don} theo mẫu pháp lý Việt Nam: quốc hiệu, tiêu ngữ, tiêu đề, kính gửi,
nội dung (có căn cứ pháp lý), kính đề nghị, cam kết. Để [TRỐNG] nơi cần điền thêm."""
    messages = [
        {"role": "user", "content": f"Thông tin vụ việc:\n{noi_dung}"},
        {"role": "user", "content": f"Soạn: {loai_don}.\n{them}"},
    ]
    return goi_claude(messages, system)

def hoi_dap(lich_su, cau_hoi, files=None):
    system = f"Bạn là chuyên gia pháp lý tại {TEN_CONG_TY}. Trả lời chuyên nghiệp, dẫn chiếu luật cụ thể khi cần."
    messages = list(lich_su)
    if files:
        for item in files:
            if item["loai"] != "anh":
                messages.insert(0, {"role": "user", "content": f"[Hồ sơ - {item['ten']}]:\n{item['du_lieu'][:2000]}"})
    messages.append({"role": "user", "content": cau_hoi})
    return goi_claude(messages, system)

# ─────────────────────────────────────────────
#  TẠO FILE WORD — ĐỊNH DẠNG MTL PREMIUM
# ─────────────────────────────────────────────
# Màu MTL chuẩn (lấy từ file mẫu công ty)
C_NAVY      = RGBColor(0x1B, 0x4A, 0x7A)   # Navy chính
C_NAVY_DARK = RGBColor(0x16, 0x3D, 0x66)   # Navy đậm
C_GOLD      = RGBColor(0xB8, 0x97, 0x3A)   # Vàng đồng
C_GOLD2     = RGBColor(0xCD, 0xB0, 0x60)   # Vàng sáng
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK      = RGBColor(0x1E, 0x29, 0x3B)   # Chữ chính
C_MUTED     = RGBColor(0x64, 0x74, 0x8B)   # Chữ phụ

def _set_cell_bg(cell, hex_color):
    """Tô màu nền ô bảng"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Đặt viền ô"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def _run(para, text, size, bold=False, italic=False, color=None):
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run

def tao_file_word(tieu_de, noi_dung, ten_ls, chuc_vu):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Inches

    doc = Document()

    # ── Thiết lập trang A4 ──
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(1.2)
    sec.bottom_margin = Cm(1.2)
    sec.left_margin   = Cm(1.2)
    sec.right_margin  = Cm(1.2)

    # Xóa spacing mặc định
    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)

    # ── HEADER: Logo + Tiêu đề tài liệu ──
    tbl_header = doc.add_table(rows=1, cols=2)
    tbl_header.style = 'Table Grid'
    tbl_header.autofit = False
    tbl_header.columns[0].width = Cm(7)
    tbl_header.columns[1].width = Cm(11.8)

    # Ô trái: logo
    cell_logo = tbl_header.cell(0, 0)
    _set_cell_bg(cell_logo, "1B4A7A")
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Nhúng logo nếu có file
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.jpg")
    if os.path.exists(logo_path):
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Cm(5.5))
    else:
        _run(p_logo, "MINHTU LAW CO., LTD", 12, bold=True, color=C_WHITE)

    # Ô phải: loại tài liệu
    cell_title = tbl_header.cell(0, 1)
    _set_cell_bg(cell_title, "FFFFFF")
    cell_title.width = Cm(11.8)
    p_type = cell_title.paragraphs[0]
    p_type.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_type.paragraph_format.space_before = Pt(12)
    _run(p_type, tieu_de.upper(), 11, bold=True, color=C_NAVY)
    p_sub = cell_title.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(p_sub, f"Ngày {datetime.now().strftime('%d/%m/%Y')}", 8.5, color=C_MUTED)
    p_sub2 = cell_title.add_paragraph()
    p_sub2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(p_sub2, f"Người soạn: {ten_ls}  ·  {chuc_vu}", 8, italic=True, color=C_MUTED)

    doc.add_paragraph()

    # ── BANNER: RIÊNG TƯ & BẢO MẬT ──
    tbl_banner = doc.add_table(rows=1, cols=1)
    tbl_banner.style = 'Table Grid'
    tbl_banner.columns[0].width = Cm(18.6)
    cell_banner = tbl_banner.cell(0, 0)
    _set_cell_bg(cell_banner, "1B4A7A")
    p_banner = cell_banner.paragraphs[0]
    p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_banner.paragraph_format.space_before = Pt(4)
    p_banner.paragraph_format.space_after  = Pt(4)
    _run(p_banner, "✦   THÔNG TIN BẢO MẬT  ·  CONFIDENTIAL   ✦", 7.5, bold=True, color=C_WHITE)

    doc.add_paragraph()

    # ── NỘI DUNG CHÍNH ──
    dong_list = [d for d in noi_dung.split("\n")]
    i = 0
    while i < len(dong_list):
        dong = dong_list[i].strip()
        dong_sach = dong.replace("**", "").replace("###", "").replace("##", "").replace("# ", "").strip()

        if not dong_sach:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Phát hiện tiêu đề section (toàn hoa, bắt đầu số, hoặc **text**)
        la_section = bool(
            (dong.isupper() and len(dong) > 3) or
            re.match(r"^\d+[\.\)]\s+[A-ZÁÀẢÃẠ]", dong) or
            (dong.startswith("**") and dong.endswith("**"))
        )

        if la_section:
            # Section header: navy nền, chữ trắng
            tbl_sec = doc.add_table(rows=1, cols=1)
            tbl_sec.style = 'Table Grid'
            tbl_sec.columns[0].width = Cm(18.6)
            cell_sec = tbl_sec.cell(0, 0)
            _set_cell_bg(cell_sec, "1B4A7A")
            p_sec = cell_sec.paragraphs[0]
            p_sec.paragraph_format.space_before = Pt(5)
            p_sec.paragraph_format.space_after  = Pt(5)
            p_sec.paragraph_format.left_indent  = Cm(0.3)
            _run(p_sec, dong_sach, 9.5, bold=True, color=C_WHITE)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        else:
            # Đoạn nội dung thường
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after      = Pt(4)
            p.paragraph_format.left_indent      = Cm(0.5)
            p.paragraph_format.first_line_indent = Cm(0)

            # Sub-heading nhỏ (bắt đầu bằng •, -, số)
            la_bullet = bool(re.match(r"^[•\-\*]\s", dong_sach) or re.match(r"^\d+\.\s", dong_sach))

            if la_bullet:
                p.paragraph_format.left_indent = Cm(1.0)
                _run(p, dong_sach, 9.5, color=C_DARK)
            else:
                _run(p, dong_sach, 9.5, color=C_DARK)

        i += 1

    # ── CHỮ KÝ ──
    doc.add_paragraph()
    tbl_sign = doc.add_table(rows=1, cols=2)
    tbl_sign.style = 'Table Grid'
    tbl_sign.columns[0].width = Cm(9.3)
    tbl_sign.columns[1].width = Cm(9.3)

    cell_l = tbl_sign.cell(0, 0)
    cell_r = tbl_sign.cell(0, 1)
    _set_cell_bg(cell_l, "F8FAFC")
    _set_cell_bg(cell_r, "F8FAFC")

    p_l = cell_l.paragraphs[0]
    p_l.paragraph_format.space_before = Pt(8)
    p_l.paragraph_format.space_after  = Pt(8)
    p_l.paragraph_format.left_indent  = Cm(0.5)
    _run(p_l, "Kính trân trọng,\n", 9.5, color=C_MUTED)
    p_l2 = cell_l.add_paragraph()
    p_l2.paragraph_format.left_indent = Cm(0.5)
    _run(p_l2, f"\n\n{ten_ls}", 10, bold=True, color=C_NAVY)
    p_l3 = cell_l.add_paragraph()
    p_l3.paragraph_format.left_indent = Cm(0.5)
    _run(p_l3, chuc_vu, 8.5, color=C_MUTED)
    p_l4 = cell_l.add_paragraph()
    p_l4.paragraph_format.left_indent = Cm(0.5)
    p_l4.paragraph_format.space_after = Pt(8)
    _run(p_l4, TEN_CONG_TY, 8.5, italic=True, color=C_GOLD)

    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r.paragraph_format.space_before = Pt(8)
    p_r.paragraph_format.right_indent = Cm(0.5)
    _run(p_r, "Hiệu lực văn bản\n", 8, bold=True, color=C_GOLD)
    p_r2 = cell_r.add_paragraph()
    p_r2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r2.paragraph_format.right_indent = Cm(0.5)
    _run(p_r2, f"TP. Hồ Chí Minh, {datetime.now().strftime('%d/%m/%Y')}", 9.5, bold=True, color=C_NAVY)
    p_r3 = cell_r.add_paragraph()
    p_r3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r3.paragraph_format.right_indent = Cm(0.5)
    p_r3.paragraph_format.space_after  = Pt(8)
    _run(p_r3, "Tài liệu soạn bởi MTL Legal Agent Premium", 7.5, italic=True, color=C_MUTED)

    # ── FOOTER: Gold bar ──
    doc.add_paragraph()
    tbl_footer = doc.add_table(rows=1, cols=1)
    tbl_footer.style = 'Table Grid'
    tbl_footer.columns[0].width = Cm(18.6)
    cell_f = tbl_footer.cell(0, 0)
    _set_cell_bg(cell_f, "B8973A")
    p_f = cell_f.paragraphs[0]
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_f.paragraph_format.space_before = Pt(3)
    p_f.paragraph_format.space_after  = Pt(3)
    _run(p_f, f"© 2026 MINHTU LAW CO., LTD  |  OUR EXPERIENCE IS YOUR SUCCESS  |  {SBT_CT}", 6.5, bold=True, color=C_WHITE)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ─────────────────────────────────────────────
#  GIAO DIỆN CHÍNH
# ─────────────────────────────────────────────
nd = st.session_state.nguoi_dung

# ── SIDEBAR ──
with st.sidebar:
    st.markdown(f"""
<div style="text-align:center;padding:16px 0 8px;">
  <div style="display:flex;align-items:center;justify-content:center;gap:3px;margin-bottom:8px;">
    <div class="mtl-box mtl-box-navy" style="width:34px;height:34px;font-size:1rem;">M</div>
    <div class="mtl-box mtl-box-gold"  style="width:34px;height:34px;font-size:1rem;">T</div>
    <div class="mtl-box mtl-box-navy" style="width:34px;height:34px;font-size:1rem;">L</div>
  </div>
  <div style="font-size:0.72rem;letter-spacing:1.5px;color:{MTL_GOLD2};text-transform:uppercase;font-weight:600;">Legal Agent Premium</div>
</div>
<div style="height:1px;background:linear-gradient(to right,transparent,{MTL_GOLD},transparent);margin:4px 0 16px;"></div>
<div style="background:rgba(255,255,255,0.06);border:1px solid rgba(168,135,74,0.3);border-radius:10px;padding:12px 14px;margin-bottom:12px;">
  <div style="font-size:0.72rem;color:{MTL_GOLD};text-transform:uppercase;letter-spacing:0.5px;margin-bottom:3px;">Người dùng</div>
  <div style="font-weight:700;font-size:0.95rem;">{nd['ho_ten']}</div>
  <div style="font-size:0.78rem;opacity:0.7;">{nd['chuc_vu']}</div>
</div>
""", unsafe_allow_html=True)

    if API_KEY:
        st.markdown("""
<div style="background:rgba(100,180,100,0.15);border:1px solid #4a9a4a;border-radius:8px;
padding:8px 12px;font-size:0.78rem;color:#90ee90;margin-bottom:12px;">
✅ API Key đã được cấu hình
</div>""", unsafe_allow_html=True)
        # Debug: hiện 8 ký tự đầu để xác nhận đúng key
        st.markdown(f"<div style='font-size:0.7rem;color:#888;margin-bottom:8px;'>Key: {API_KEY[:12]}...{API_KEY[-4:]}</div>", unsafe_allow_html=True)
    else:
        st.markdown("""
<div style="background:rgba(220,80,80,0.15);border:1px solid #c04040;border-radius:8px;
padding:8px 12px;font-size:0.78rem;color:#ff9090;margin-bottom:12px;">
⚠️ Chưa có API Key!<br>Mở file app.py → tìm dòng ANTHROPIC_API_KEY → dán key vào
</div>""", unsafe_allow_html=True)

    st.markdown(f"<div style='height:1px;background:rgba(168,135,74,0.25);margin:4px 0 12px;'></div>", unsafe_allow_html=True)
    st.markdown(f"<div style='font-size:0.8rem;color:{MTL_GOLD2};font-weight:600;text-transform:uppercase;letter-spacing:0.5px;margin-bottom:8px;'>📂 Tải hồ sơ lên</div>", unsafe_allow_html=True)

    files_upload = st.file_uploader(
        "Chọn file",
        type=["pdf", "docx", "png", "jpg", "jpeg", "tiff", "bmp"],
        accept_multiple_files=True,
        label_visibility="collapsed",
    )

    if "noi_dung_files" not in st.session_state:
        st.session_state.noi_dung_files = []

    if files_upload:
        st.session_state.noi_dung_files = []
        for f in files_upload:
            data = f.read()
            ext  = f.name.rsplit(".", 1)[-1].lower()
            if ext == "pdf":
                st.session_state.noi_dung_files.append({"ten": f.name, "loai": "pdf", "du_lieu": doc_pdf(data)})
            elif ext == "docx":
                st.session_state.noi_dung_files.append({"ten": f.name, "loai": "docx", "du_lieu": doc_docx(data)})
            elif ext in ["png", "jpg", "jpeg", "tiff", "bmp"]:
                media = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png", "tiff": "image/tiff", "bmp": "image/bmp"}
                st.session_state.noi_dung_files.append({
                    "ten": f.name, "loai": "anh",
                    "du_lieu": base64.standard_b64encode(data).decode(),
                    "media_type": media.get(ext, "image/jpeg"),
                })
        if st.session_state.noi_dung_files:
            st.success(f"✅ Đã tải {len(st.session_state.noi_dung_files)} file")
            for item in st.session_state.noi_dung_files:
                icon = "🖼️" if item["loai"] == "anh" else "📄"
                st.markdown(f"<div style='font-size:0.8rem;padding:2px 0;'>{icon} {item['ten']}</div>", unsafe_allow_html=True)

    st.markdown(f"<div style='height:1px;background:rgba(168,135,74,0.25);margin:12px 0;'></div>", unsafe_allow_html=True)

    # ── Giá trị cốt lõi ở sidebar ──
    st.markdown(f"""
<div style="margin-bottom:12px;">
  <div style="font-size:0.62rem;color:{MTL_GOLD};letter-spacing:2px;font-weight:700;
  text-transform:uppercase;text-align:center;margin-bottom:8px;">Giá trị cốt lõi</div>
  <div style="display:flex;flex-direction:column;gap:5px;">
    <div style="display:flex;align-items:center;gap:8px;background:rgba(168,135,74,0.1);
    border-left:2px solid {MTL_GOLD};border-radius:0 6px 6px 0;padding:6px 10px;">
      <span style="font-size:0.9rem;">🤝</span>
      <div>
        <div style="font-size:0.7rem;font-weight:700;color:{MTL_GOLD};letter-spacing:0.5px;">CAM KẾT</div>
        <div style="font-size:0.62rem;color:rgba(232,238,245,0.5);margin-top:1px;">Tận tâm phục vụ đến cùng</div>
      </div>
    </div>
    <div style="display:flex;align-items:center;gap:8px;background:rgba(168,135,74,0.1);
    border-left:2px solid {MTL_GOLD};border-radius:0 6px 6px 0;padding:6px 10px;">
      <span style="font-size:0.9rem;">⚖️</span>
      <div>
        <div style="font-size:0.7rem;font-weight:700;color:{MTL_GOLD};letter-spacing:0.5px;">CHÍNH TRỰC</div>
        <div style="font-size:0.62rem;color:rgba(232,238,245,0.5);margin-top:1px;">Minh bạch & đạo đức nghề nghiệp</div>
      </div>
    </div>
    <div style="display:flex;align-items:center;gap:8px;background:rgba(168,135,74,0.1);
    border-left:2px solid {MTL_GOLD};border-radius:0 6px 6px 0;padding:6px 10px;">
      <span style="font-size:0.9rem;">📚</span>
      <div>
        <div style="font-size:0.7rem;font-weight:700;color:{MTL_GOLD};letter-spacing:0.5px;">HỌC HỎI</div>
        <div style="font-size:0.62rem;color:rgba(232,238,245,0.5);margin-top:1px;">Không ngừng trau dồi kiến thức</div>
      </div>
    </div>
  </div>
</div>
""", unsafe_allow_html=True)

    if st.button("🚪 Đăng xuất", use_container_width=True):
        dang_xuat()

    # ── GMAIL OAuth ────────────────────────────────────
    st.markdown(
        f"<div style='height:1px;background:rgba(168,135,74,0.25);margin:12px 0;'></div>",
        unsafe_allow_html=True,
    )
    st.markdown(
        f"<div style='font-size:0.8rem;color:{MTL_GOLD2};font-weight:600;"
        f"text-transform:uppercase;letter-spacing:0.5px;margin-bottom:8px;'>📧 Gmail</div>",
        unsafe_allow_html=True,
    )

    import requests as _rq

    _CLIENT_ID     = os.environ.get("GOOGLE_CLIENT_ID", "").strip()
    _CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET", "").strip()
    _REDIRECT_URI  = "https://web-production-57eec.up.railway.app"
    _SCOPES = [
        "https://www.googleapis.com/auth/gmail.readonly",
        "https://www.googleapis.com/auth/gmail.send",
    ]
    _cred_key = f"gcred_{nd['ten_tk']}"

    # Cache đảm bảo mỗi code chỉ được exchange ĐÚNG 1 LẦN
    # dù Streamlit chạy bao nhiêu thread/rerun đi nữa
    @st.cache_data(ttl=120, show_spinner=False)
    def _do_exchange(code, client_id, client_secret, redirect_uri):
        import requests as _r
        resp = _r.post(
            "https://oauth2.googleapis.com/token",
            data={
                "code":          code,
                "client_id":     client_id,
                "client_secret": client_secret,
                "redirect_uri":  redirect_uri,
                "grant_type":    "authorization_code",
            },
            timeout=15,
        )
        return resp.json()

    def _get_email(access_token):
        try:
            return _rq.get(
                "https://www.googleapis.com/oauth2/v2/userinfo",
                headers={"Authorization": f"Bearer {access_token}"},
                timeout=8,
            ).json().get("email", "")
        except Exception:
            return ""

    def _is_connected():
        c = st.session_state.get(_cred_key)
        if not c:
            return False
        if not c.valid:
            try:
                c.refresh(Request())
                st.session_state[_cred_key] = c
            except Exception:
                return False
        return c.valid

    if not _CLIENT_ID or not _CLIENT_SECRET:
        st.error("Thiếu GOOGLE_CLIENT_ID hoặc CLIENT_SECRET")

    elif _is_connected():
        _em = st.session_state.get(f"gemail_{nd['ten_tk']}", "Gmail")
        st.markdown(
            f"<div style='background:rgba(100,180,100,0.15);border:1px solid #4a9a4a;"
            f"border-radius:8px;padding:8px 10px;font-size:0.78rem;color:#90ee90;'>"
            f"✅ {_em}</div>",
            unsafe_allow_html=True,
        )
        if st.button("↩ Ngắt kết nối", use_container_width=True,
                     key=f"gdisconn_{nd['ten_tk']}"):
            st.session_state.pop(_cred_key, None)
            st.session_state.pop(f"gemail_{nd['ten_tk']}", None)
            st.rerun()

    else:
        _qp    = st.query_params
        _code  = _qp.get("code", "")
        _state = _qp.get("state", "")

        if _code and _state == nd["ten_tk"]:
            # Xoá URL ngay để tránh các rerun sau thấy lại code
            st.query_params.clear()

            # Gọi hàm đã được cache — dù gọi nhiều lần cũng chỉ POST 1 lần
            _tok = _do_exchange(_code, _CLIENT_ID, _CLIENT_SECRET, _REDIRECT_URI)

            if "access_token" in _tok:
                _creds = Credentials(
                    token=_tok["access_token"],
                    refresh_token=_tok.get("refresh_token"),
                    token_uri="https://oauth2.googleapis.com/token",
                    client_id=_CLIENT_ID,
                    client_secret=_CLIENT_SECRET,
                    scopes=_SCOPES,
                )
                st.session_state[_cred_key] = _creds
                st.session_state[f"gemail_{nd['ten_tk']}"] = _get_email(_tok["access_token"])
                st.rerun()
            else:
                _err = _tok.get("error_description", _tok.get("error", "?"))
                st.error(f"❌ Lỗi: {_err}")

        else:
            # Nút đăng nhập
            import urllib.parse as _up
            _auth_url = (
                "https://accounts.google.com/o/oauth2/v2/auth?"
                + _up.urlencode({
                    "client_id":     _CLIENT_ID,
                    "redirect_uri":  _REDIRECT_URI,
                    "response_type": "code",
                    "scope":         " ".join(_SCOPES),
                    "access_type":   "offline",
                    "prompt":        "consent",
                    "state":         nd["ten_tk"],
                })
            )
            st.markdown(
                f'<a href="{_auth_url}" target="_self">'
                f'<button style="width:100%;background:#4285F4;color:white;'
                f'border:none;border-radius:8px;padding:9px;font-size:0.85rem;'
                f'font-weight:600;cursor:pointer;">'
                f'🔐 Đăng nhập với Google</button></a>',
                unsafe_allow_html=True,
            )
            st.caption("Dùng mật khẩu Gmail thông thường")
# ── HEADER ──
st.markdown(f"""
<div class="mtl-header">
  <div class="mtl-header-inner">
    <div style="display:flex;align-items:center;gap:4px;flex-shrink:0;">
      <div class="mtl-box mtl-box-navy">M</div>
      <div class="mtl-box mtl-box-gold">T</div>
      <div class="mtl-box mtl-box-navy">L</div>
    </div>
    <div class="mtl-divider"></div>
    <div class="mtl-title-block">
      <h1>LUẬT MINH TÚ</h1>
      <p class="mtl-sub">⚖ Legal Agent Premium &nbsp;·&nbsp; Hệ thống hỗ trợ pháp lý thông minh</p>
    </div>
    <div class="mtl-user-badge">
      <div class="name">{nd['ho_ten']}</div>
      <div class="role">{nd['chuc_vu']}</div>
      <div class="date">{datetime.now().strftime('%d/%m/%Y  %H:%M')}</div>
    </div>
  </div>
</div>

<!-- Thanh giá trị cốt lõi -->
<div style="background:linear-gradient(90deg,{MTL_NAVY2} 0%,#122d50 100%);
border-bottom:2px solid {MTL_GOLD}44;
padding:7px 28px;margin-top:-4px;margin-bottom:8px;
display:flex;align-items:center;justify-content:center;gap:0;">

  <div style="display:flex;align-items:center;gap:8px;padding:0 24px;
  border-right:1px solid {MTL_GOLD}40;">
    <span style="font-size:1rem;">🤝</span>
    <div>
      <div style="font-size:0.7rem;font-weight:800;color:{MTL_GOLD};
      letter-spacing:1.5px;text-transform:uppercase;line-height:1;">Cam kết</div>
      <div style="font-size:0.6rem;color:rgba(201,169,110,0.6);margin-top:1px;">Tận tâm phục vụ đến cùng</div>
    </div>
  </div>

  <div style="display:flex;align-items:center;gap:8px;padding:0 24px;
  border-right:1px solid {MTL_GOLD}40;">
    <span style="font-size:1rem;">⚖️</span>
    <div>
      <div style="font-size:0.7rem;font-weight:800;color:{MTL_GOLD};
      letter-spacing:1.5px;text-transform:uppercase;line-height:1;">Chính trực</div>
      <div style="font-size:0.6rem;color:rgba(201,169,110,0.6);margin-top:1px;">Minh bạch & đạo đức nghề nghiệp</div>
    </div>
  </div>

  <div style="display:flex;align-items:center;gap:8px;padding:0 24px;">
    <span style="font-size:1rem;">📚</span>
    <div>
      <div style="font-size:0.7rem;font-weight:800;color:{MTL_GOLD};
      letter-spacing:1.5px;text-transform:uppercase;line-height:1;">Học hỏi</div>
      <div style="font-size:0.6rem;color:rgba(201,169,110,0.6);margin-top:1px;">Không ngừng trau dồi kiến thức</div>
    </div>
  </div>

  <div style="margin-left:auto;font-size:0.6rem;color:rgba(201,169,110,0.4);
  font-style:italic;white-space:nowrap;">OUR EXPERIENCE IS YOUR SUCCESS</div>
</div>
""", unsafe_allow_html=True)

# ── TABS ──
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "🔍 Phân tích hồ sơ",
    "📝 Soạn thảo văn bản",
    "💬 Hỏi đáp pháp lý",
    "📋 Hướng dẫn sử dụng",
    "📧 Email Intelligence",
    "📌 Quản lý Công việc",
])

# ══════════════════════════════════════════════
#  TAB 1 — PHÂN TÍCH HỒ SƠ
# ══════════════════════════════════════════════
with tab1:
    st.subheader("🔍 Phân tích hồ sơ vụ việc")
    col_a, col_b = st.columns([2, 1])
    with col_a:
        yeu_cau = st.text_area(
            "Yêu cầu phân tích cụ thể (tùy chọn)",
            placeholder="Ví dụ: Tập trung vào thời hiệu khởi kiện, quyền đòi bồi thường...",
            height=80,
        )
    with col_b:
        st.markdown("<br>", unsafe_allow_html=True)
        nut_pt = st.button("🚀 Phân tích ngay", use_container_width=True, type="primary")

    if nut_pt:
        if not API_KEY:
            st.error("❌ Chưa có API Key. Mở file app.py, tìm dòng ANTHROPIC_API_KEY và điền key vào.")
        elif not st.session_state.noi_dung_files:
            st.warning("⚠️ Vui lòng tải ít nhất 1 file hồ sơ ở thanh bên trái.")
        else:
            with st.spinner("🤖 AI đang nghiên cứu hồ sơ..."):
                ket_qua = phan_tich_ho_so(st.session_state.noi_dung_files, yeu_cau)
                st.session_state.ket_qua_phan_tich = ket_qua

    if "ket_qua_phan_tich" in st.session_state and st.session_state.ket_qua_phan_tich:
        st.markdown("---")
        st.markdown("#### 📊 Kết quả phân tích")
        st.markdown(
            f'<div class="result-box">{st.session_state.ket_qua_phan_tich.replace(chr(10), "<br>")}</div>',
            unsafe_allow_html=True,
        )
        st.markdown("<br>", unsafe_allow_html=True)
        word_bytes = tao_file_word(
            "BÁO CÁO PHÂN TÍCH HỒ SƠ VỤ VIỆC",
            st.session_state.ket_qua_phan_tich,
            nd["ho_ten"], nd["chuc_vu"],
        )
        st.download_button(
            "⬇️ Tải xuống file Word", data=word_bytes,
            file_name=f"PhanTich_{datetime.now().strftime('%d%m%Y_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

# ══════════════════════════════════════════════
#  TAB 2 — SOẠN THẢO
# ══════════════════════════════════════════════
with tab2:
    st.subheader("📝 Soạn thảo văn bản pháp lý")

    LOAI_DON = [
        "Đơn khởi kiện", "Đơn yêu cầu thi hành án", "Đơn đề nghị hòa giải",
        "Đơn tố cáo", "Đơn xin cấp bản sao hồ sơ",
        "Đơn xin gia hạn nộp tiền tạm ứng án phí",
        "Hợp đồng dịch vụ pháp lý", "Thông báo pháp lý (Legal Notice)",
        "Biên bản cuộc họp", "Phiếu yêu cầu tư vấn", "Văn bản khác (tự nhập)",
    ]

    col1, col2 = st.columns(2)
    with col1:
        loai_don = st.selectbox("Loại văn bản", LOAI_DON)
    with col2:
        if loai_don == "Văn bản khác (tự nhập)":
            loai_don = st.text_input("Nhập loại văn bản", placeholder="Ví dụ: Đơn phản đối...")

    noi_dung_vv = st.text_area(
        "Mô tả vụ việc / thông tin cần đưa vào",
        placeholder="Nguyên đơn: ...\nBị đơn: ...\nNội dung tranh chấp: ...\nYêu cầu: ...",
        height=140,
    )

    if "ket_qua_phan_tich" in st.session_state and st.session_state.ket_qua_phan_tich:
        if st.checkbox("📂 Lấy thông tin từ hồ sơ đã phân tích"):
            noi_dung_vv = st.session_state.ket_qua_phan_tich[:1500]

    them = st.text_input("Yêu cầu thêm", placeholder="Ví dụ: Nhấn mạnh điều 166 BLDS...")

    if st.button("✍️ Soạn văn bản", type="primary"):
        if not API_KEY:
            st.error("❌ Chưa có API Key. Mở file app.py, tìm dòng ANTHROPIC_API_KEY và điền key vào.")
        elif not noi_dung_vv.strip():
            st.warning("⚠️ Vui lòng nhập thông tin vụ việc.")
        else:
            with st.spinner("🤖 AI đang soạn thảo..."):
                van_ban = soan_don_tu(loai_don, noi_dung_vv, them)
                st.session_state.van_ban_soan = van_ban
                st.session_state.loai_van_ban = loai_don

    if "van_ban_soan" in st.session_state and st.session_state.van_ban_soan:
        st.markdown("---")
        vb_edit = st.text_area("✏️ Chỉnh sửa nếu cần", value=st.session_state.van_ban_soan, height=400)
        ten_file = st.session_state.loai_van_ban.replace(" ", "_")
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            word_bytes = tao_file_word(st.session_state.loai_van_ban, vb_edit, nd["ho_ten"], nd["chuc_vu"])
            st.download_button(
                "⬇️ Tải xuống file Word", data=word_bytes,
                file_name=f"{ten_file}_{datetime.now().strftime('%d%m%Y')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with col_d2:
            st.download_button(
                "📋 Tải xuống file TXT", data=vb_edit.encode("utf-8"),
                file_name=f"{ten_file}_{datetime.now().strftime('%d%m%Y')}.txt",
                mime="text/plain", use_container_width=True,
            )

# ══════════════════════════════════════════════
#  TAB 3 — HỎI ĐÁP
# ══════════════════════════════════════════════
with tab3:
    st.subheader("💬 Hỏi đáp pháp lý")

    if "lich_su_chat" not in st.session_state:
        st.session_state.lich_su_chat = []

    for tin in st.session_state.lich_su_chat:
        with st.chat_message("user" if tin["role"] == "user" else "assistant",
                             avatar="👤" if tin["role"] == "user" else "⚖️"):
            st.write(tin["content"])

    cau_hoi = st.chat_input("Hỏi về pháp luật, vụ việc, hoặc nội dung hồ sơ đã tải lên...")

    if cau_hoi:
        if not API_KEY:
            st.error("❌ Chưa có API Key. Mở file app.py, tìm dòng ANTHROPIC_API_KEY và điền key vào.")
        else:
            st.session_state.lich_su_chat.append({"role": "user", "content": cau_hoi})
            with st.chat_message("user", avatar="👤"):
                st.write(cau_hoi)
            with st.chat_message("assistant", avatar="⚖️"):
                with st.spinner("Đang tra cứu..."):
                    tra_loi = hoi_dap(st.session_state.lich_su_chat[:-1], cau_hoi, st.session_state.noi_dung_files)
                    st.write(tra_loi)
                    st.session_state.lich_su_chat.append({"role": "assistant", "content": tra_loi})

    if st.session_state.lich_su_chat:
        if st.button("🗑️ Xóa lịch sử chat"):
            st.session_state.lich_su_chat = []
            st.rerun()

# ══════════════════════════════════════════════
#  TAB 4 — HƯỚNG DẪN
# ══════════════════════════════════════════════
with tab4:
    st.subheader("📋 Hướng dẫn sử dụng MTL Legal Agent Premium")
    st.markdown(f"""
### 🔑 Cấu hình API Key (chỉ làm 1 lần)

Mở file `app.py` → tìm dòng số **21**:
```
ANTHROPIC_API_KEY = "sk-ant-"
```
Thay `sk-ant-` bằng key thật của bạn. Lấy key tại: **console.anthropic.com**

---

### 🚀 Sử dụng hàng ngày

**Bước 1:** Tải hồ sơ vụ việc ở thanh bên trái (PDF, Word, ảnh chụp, chữ viết tay).

**Bước 2:** Chọn chức năng:
- **Phân tích hồ sơ** — AI đọc toàn bộ hồ sơ, phân tích pháp lý, xuất báo cáo Word.
- **Soạn thảo văn bản** — Chọn loại đơn, AI soạn đúng mẫu pháp lý, tải về Word.
- **Hỏi đáp pháp lý** — Chat trực tiếp về luật và hồ sơ đã tải lên.

---

### 👥 Danh sách tài khoản

| Tài khoản | Mật khẩu | Họ tên |
|-----------|----------|--------|
""" + "\n".join([f"| `{tk}` | `{info['mat_khau']}` | {info['ho_ten']} |" for tk, info in TAI_KHOAN.items()]) + f"""

---
### 🏢 {TEN_CONG_TY}
{DIA_CHI_CT}  
{DIA_CHI_DN}  
{SBT_CT}
""")

# ══════════════════════════════════════════════
#  TAB 5 — EMAIL INTELLIGENCE
# ══════════════════════════════════════════════

# ── Gmail API helpers ──────────────────────────────────────────
def _gmail_service():
    """Trả về Gmail service của luật sư đang đăng nhập."""
    creds = st.session_state.get(f"gcred_{nd['ten_tk']}")
    if not creds or not creds.valid:
        return None
    return build("gmail", "v1", credentials=creds)


def tai_email_gmail_api(so_luong: int = 12) -> list:
    """Tải email từ Gmail qua API."""
    svc = _gmail_service()
    if not svc:
        return []
    try:
        resp = svc.users().messages().list(
            userId="me", maxResults=so_luong, labelIds=["INBOX"]
        ).execute()
        msgs = resp.get("messages", [])
        results = []
        for m in msgs:
            full = svc.users().messages().get(
                userId="me", id=m["id"], format="full"
            ).execute()
            headers = {h["name"]: h["value"] for h in full["payload"]["headers"]}
            subject = headers.get("Subject", "(Không có tiêu đề)")
            from_   = headers.get("From", "")
            date_   = headers.get("Date", "")[:22]

            # Parse sender
            _m = re.match(r'"?(.+?)"?\s*<(.+?)>', from_)
            from_name  = _m.group(1).strip() if _m else from_
            from_email = _m.group(2).strip() if _m else from_

            # Lấy text body
            body = _extract_body(full["payload"])

            unread = "UNREAD" in full.get("labelIds", [])
            results.append({
                "id":        m["id"],
                "fromName":  from_name,
                "fromEmail": from_email,
                "subject":   subject,
                "date":      date_,
                "body":      body[:4000],
                "unread":    unread,
            })
        return results
    except Exception as e:
        st.error(f"Gmail API lỗi: {e}")
        return []


def _extract_body(payload: dict) -> str:
    """Lấy nội dung text/plain từ payload email."""
    if payload.get("mimeType") == "text/plain":
        data = payload.get("body", {}).get("data", "")
        if data:
            return base64.urlsafe_b64decode(data).decode("utf-8", errors="ignore")
    for part in payload.get("parts", []):
        result = _extract_body(part)
        if result:
            return result
    return ""


def gui_email_gmail_api(to: str, subject: str, body: str) -> bool:
    """Gửi email reply qua Gmail API."""
    svc = _gmail_service()
    if not svc:
        return False
    try:
        import email as email_lib
        from email.mime.text import MIMEText
        msg = MIMEText(body, "plain", "utf-8")
        sender = st.session_state.get(f"gemail_{nd['ten_tk']}", "me")
        msg["From"]    = sender
        msg["To"]      = to
        msg["Subject"] = f"Re: {subject}"
        raw = base64.urlsafe_b64encode(msg.as_bytes()).decode()
        svc.users().messages().send(
            userId="me", body={"raw": raw}
        ).execute()
        return True
    except Exception as e:
        st.error(f"Gửi email lỗi: {e}")
        return False


# ── AI helpers — tái sử dụng goi_claude() ──────────────────────
def phan_tich_email_phap_ly(email: dict) -> dict:
    system = f"Bạn là trợ lý pháp lý tại {TEN_CONG_TY}. Trả về JSON thuần, không markdown."
    prompt = f"""Phân tích email và trả về đúng JSON:
{{
  "urgency": "high|medium|low", "urgency_score": 0-100,
  "urgency_reason": "lý do", "category": "loại vụ việc",
  "summary": "tóm tắt 1-2 câu",
  "legal_issues": ["vấn đề 1","vấn đề 2"],
  "relevant_laws": ["Luật 1","Luật 2"],
  "parties": [{{"role":"vai trò","name":"tên"}}],
  "action_items": ["việc 1","việc 2"],
  "deadline": "thời hạn hoặc null",
  "risk_level": "Cao|Trung bình|Thấp"
}}
Tiêu đề: {email.get('subject','')}
Từ: {email.get('fromName','')} <{email.get('fromEmail','')}>
Nội dung:\n{email.get('body','')[:3000]}"""
    text = goi_claude([{"role":"user","content":prompt}], system)
    try:
        return json.loads(text.replace("```json","").replace("```","").strip())
    except Exception:
        return {}


def soan_phan_hoi(email: dict, analysis: dict, tone: str) -> str:
    tone_map = {
        "formal":   "trang trọng, văn phong luật sư chuyên nghiệp",
        "friendly": "thân thiện, gần gũi nhưng vẫn chuyên nghiệp",
        "firm":     "kiên quyết, rõ ràng, thể hiện quyền hạn",
        "urgent":   "khẩn cấp, nhấn mạnh cần hành động ngay",
    }
    ctx = ""
    if analysis:
        ctx = (f"\nPhân tích: {analysis.get('summary','')}\n"
               f"Hành động: {'; '.join(analysis.get('action_items',[]))}")
    system = f"Bạn là luật sư tại {TEN_CONG_TY}."
    prompt = (f"Soạn email phản hồi tiếng Việt, giọng {tone_map.get(tone,'trang trọng')}. "
              f"Bắt đầu 'Kính gửi...', KHÔNG viết subject, xác nhận nhận email, "
              f"nêu hướng xử lý, đề xuất bước tiếp theo. "
              f"Ký tên: {nd['ho_ten']} — {TEN_CONG_TY}{ctx}\n\n"
              f"Tiêu đề: {email.get('subject','')}\n"
              f"Từ: {email.get('fromName','')}\n"
              f"Nội dung:\n{email.get('body','')[:2000]}")
    return goi_claude([{"role":"user","content":prompt}], system)


def gan_tag(email: dict) -> list:
    text = f"{email.get('subject','')} {email.get('body','')}".lower()
    rules = {
        "🔴 Khẩn":       ["khẩn","gấp","ngay","vi phạm","khởi kiện"],
        "🟡 Hợp đồng":   ["hợp đồng","ký kết","điều khoản","soát xét"],
        "🟣 Tranh chấp": ["tranh chấp","kiện","tòa án","bồi thường"],
        "🟢 Tư vấn":     ["tư vấn","hỏi","thành lập","startup"],
    }
    tags = [t for t, kws in rules.items() if any(k in text for k in kws)]
    return tags if tags else ["🔵 Thông thường"]


EMAIL_MAU = [
    {"id":"m1","unread":True,"fromName":"Nguyễn Văn Minh","fromEmail":"nvminh@vietcorp.vn",
     "date":"09:42","subject":"Tranh chấp hợp đồng mua bán căn hộ — cần tư vấn khẩn",
     "body":"Kính gửi Luật sư,\n\nTôi đã ký hợp đồng mua căn hộ tại dự án Green Valley ngày 15/03/2024, giá trị 3,2 tỷ đồng. Chủ đầu tư vi phạm:\n1. Trễ bàn giao 8 tháng\n2. Từ chối trả phạt điều 9 (0.05%/ngày)\n3. Đơn phương thay đổi thiết kế\n\nCần tư vấn khẩn.\n\nTrân trọng,\nNguyễn Văn Minh"},
    {"id":"m2","unread":True,"fromName":"Trần Thị Hà","fromEmail":"ttha@mfg.com.vn",
     "date":"Hôm qua","subject":"Soát xét hợp đồng phân phối độc quyền 5M USD",
     "body":"Luật sư kính mến,\n\nChuẩn bị ký hợp đồng phân phối với Korea Tech Co., Ltd., giá trị 5M USD/năm. Cần soát xét Điều 6, 12, 15 và Phụ lục A.\nHạn ký: 30/04/2025.\n\nTrân trọng,\nTrần Thị Hà"},
    {"id":"m3","unread":False,"fromName":"Phạm Quốc Bảo","fromEmail":"pqbao@startup.io",
     "date":"20/04","subject":"Tư vấn thành lập startup FinTech P2P Lending",
     "body":"Kính gửi Văn phòng Luật Minh Tú,\n\nCần tư vấn thành lập startup FinTech (P2P Lending):\n1. Hình thức pháp nhân\n2. Cấu trúc vốn Seed từ Singapore\n3. NĐ 52/2021\n\nNgân sách: 50-80tr.\n\nTrân trọng, Phạm Quốc Bảo"},
]


# ══════════════════════════════════════════════
#  RENDER TAB 5
# ══════════════════════════════════════════════
with tab5:
    for _k, _v in {
        "ei_emails":[], "ei_selected":None,
        "ei_analysis":None, "ei_draft":"",
        "ei_tone":"formal", "ei_sent":[],
    }.items():
        if _k not in st.session_state:
            st.session_state[_k] = _v

    # Kiểm tra kết nối
    _connected = bool(
        st.session_state.get(f"gcred_{nd['ten_tk']}") and
        st.session_state.get(f"gcred_{nd['ten_tk']}").valid
    )
    _gmail_addr = st.session_state.get(f"gemail_{nd['ten_tk']}", "")

    # Banner
    _status = (f"<span style='color:#90ee90;'>✅ {_gmail_addr}</span>"
               if _connected else
               "<span style='color:#ffa07a;'>⚠️ Chưa kết nối — nhấn nút Google ở thanh bên</span>")
    st.markdown(f"""
<div style="background:linear-gradient(135deg,{MTL_NAVY2} 0%,{MTL_NAVY} 100%);
border-radius:10px;padding:14px 20px;margin-bottom:18px;
border-left:4px solid {MTL_GOLD};display:flex;align-items:center;justify-content:space-between;">
  <div>
    <span style="color:white;font-size:1.05rem;font-weight:700;">📧 Email Intelligence</span>
    <span style="color:{MTL_GOLD2};font-size:0.8rem;margin-left:12px;">
      Gmail API · Phân tích pháp lý AI · Soạn thảo tự động
    </span>
  </div>
  <div>{_status}</div>
</div>""", unsafe_allow_html=True)

    col_inbox, col_email, col_ai = st.columns([1.2, 2, 1.8])

    # ── CỘT 1: HỘP THƯ ──────────────────────
    with col_inbox:
        st.markdown(f"<div style='font-weight:700;color:{MTL_NAVY};margin-bottom:8px;'>📬 Hộp thư</div>",
                    unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            if st.button("↻ Tải email", use_container_width=True, key="ei_load"):
                if not _connected:
                    st.warning("Đăng nhập Google ở thanh bên trước")
                else:
                    with st.spinner("Đang tải..."):
                        emails = tai_email_gmail_api(so_luong=12)
                    if emails:
                        st.session_state.ei_emails = emails
                        st.rerun()
                    else:
                        st.error("Không tải được email")
        with c2:
            if st.button("📋 Demo", use_container_width=True, key="ei_demo"):
                st.session_state.ei_emails = EMAIL_MAU
                st.rerun()

        emails = st.session_state.ei_emails
        if not emails:
            st.markdown("<div style='color:#aaa;font-size:0.82rem;text-align:center;"
                        "padding:24px 0;'>Đăng nhập Gmail → Tải email<br>hoặc nhấn Demo</div>",
                        unsafe_allow_html=True)
        else:
            chua_doc = sum(1 for e in emails if e.get("unread"))
            st.caption(f"{chua_doc} chưa đọc · {len(emails)} tổng")
            for em in emails:
                _sel = st.session_state.ei_selected
                is_sel = bool(_sel and _sel.get("id") == em["id"])
                tags   = gan_tag(em)
                label  = ("🔵 " if em.get("unread") else "") + em["fromName"]
                subj   = em["subject"][:36] + ("…" if len(em["subject"]) > 36 else "")
                if st.button(f"{label}\n{' '.join(tags[:1])}  {subj}",
                             key=f"ei_em_{em['id']}", use_container_width=True,
                             type="primary" if is_sel else "secondary"):
                    st.session_state.ei_selected = em
                    st.session_state.ei_analysis  = None
                    st.session_state.ei_draft     = ""
                    st.rerun()

    # ── CỘT 2: NỘI DUNG EMAIL ───────────────
    with col_email:
        em = st.session_state.ei_selected
        if em is None:
            st.markdown("<div style='color:#aaa;text-align:center;padding:80px 0;'>"
                        "👈 Chọn email để xem</div>", unsafe_allow_html=True)
        else:
            st.markdown(
                f"<div style='font-size:1rem;font-weight:700;color:{MTL_NAVY};"
                f"border-bottom:2px solid {MTL_GOLD}44;padding-bottom:8px;margin-bottom:10px;'>"
                f"{em['subject']}</div>", unsafe_allow_html=True)
            m1, m2 = st.columns(2)
            m1.markdown(f"**Từ:** {em['fromName']}  \n`{em['fromEmail']}`")
            m2.markdown(f"**Lúc:** {em.get('date','')}")
            tags_html = " &nbsp;".join(
                f"<span style='background:{MTL_NAVY}11;border:1px solid {MTL_NAVY}33;"
                f"border-radius:4px;padding:2px 8px;font-size:0.75rem;'>{t}</span>"
                for t in gan_tag(em))
            st.markdown(tags_html, unsafe_allow_html=True)
            st.divider()
            body_safe = em["body"].replace("<","&lt;").replace(">","&gt;")
            st.markdown(
                f"<div style='background:#f8f9fc;border:1px solid #e0e8f5;"
                f"border-left:3px solid {MTL_NAVY};border-radius:0 8px 8px 0;"
                f"padding:16px 18px;font-size:0.87rem;line-height:1.85;"
                f"white-space:pre-wrap;max-height:360px;overflow-y:auto;'>"
                f"{body_safe}</div>", unsafe_allow_html=True)
            st.divider()
            qa, qb, qc = st.columns(3)
            with qa:
                if st.button("🔍 Phân tích AI", use_container_width=True, key="ei_analyze"):
                    with st.spinner("Claude đang phân tích..."):
                        st.session_state.ei_analysis = phan_tich_email_phap_ly(em)
                    st.rerun()
            with qb:
                if st.button("✦ Soạn thảo", use_container_width=True, key="ei_draft_btn"):
                    with st.spinner("Claude đang soạn..."):
                        st.session_state.ei_draft = soan_phan_hoi(
                            em, st.session_state.ei_analysis, st.session_state.ei_tone)
                    st.rerun()
            with qc:
                if st.button("📄 Tạo văn bản", use_container_width=True, key="ei_docbtn"):
                    a = st.session_state.ei_analysis
                    if a:
                        nd_vb = (f"Vụ việc: {em['subject']}\nKhách hàng: {em['fromName']}\n\n"
                                 f"Tóm tắt: {a.get('summary','')}\n\n"
                                 f"Vấn đề pháp lý:\n" + "\n".join(f"- {i}" for i in a.get("legal_issues",[])) +
                                 f"\n\nHành động:\n" + "\n".join(f"{i+1}. {x}" for i,x in enumerate(a.get("action_items",[]))))
                        vb = soan_don_tu("Thư tư vấn pháp lý", nd_vb)
                        st.session_state.van_ban_soan  = vb
                        st.session_state.loai_van_ban = "Thư tư vấn pháp lý"
                        st.success("✅ Đã tạo — xem tab Soạn thảo văn bản")
                    else:
                        st.warning("Phân tích AI trước")

    # ── CỘT 3: AI PANEL ─────────────────────
    with col_ai:
        if st.session_state.ei_selected is None:
            st.info("Chọn email để bắt đầu")
        else:
            em = st.session_state.ei_selected
            ai1, ai2, ai3 = st.tabs(["🔍 Phân tích", "✍ Soạn thảo", "📤 Đã gửi"])

            with ai1:
                a = st.session_state.ei_analysis
                if a is None:
                    st.markdown("<div style='color:#aaa;text-align:center;padding:24px 0;'>"
                                "Nhấn 🔍 Phân tích AI</div>", unsafe_allow_html=True)
                elif a == {}:
                    st.error("Phân tích thất bại — kiểm tra API Key")
                else:
                    score = a.get("urgency_score", 0)
                    level = a.get("urgency","low")
                    bar_c = {"high":"#e53e3e","medium":"#d69e2e","low":"#38a169"}.get(level,"#718096")
                    urg_l = {"high":"🔴 Khẩn cấp","medium":"🟡 Trung bình","low":"🟢 Thấp"}.get(level,"")
                    st.markdown(
                        f"<div style='display:flex;justify-content:space-between;margin-bottom:4px;'>"
                        f"<b style='font-size:0.85rem;'>{urg_l}</b>"
                        f"<span style='color:#718096;font-size:0.8rem;'>{score}/100</span></div>"
                        f"<div style='background:#e2e8f0;border-radius:4px;height:6px;'>"
                        f"<div style='background:{bar_c};width:{score}%;height:6px;border-radius:4px;'></div></div>",
                        unsafe_allow_html=True)
                    st.caption(a.get("urgency_reason",""))
                    if a.get("deadline"):
                        st.warning(f"⏱ {a['deadline']}")
                    st.markdown(
                        f"<div style='background:{MTL_NAVY}08;border-left:3px solid {MTL_GOLD};"
                        f"border-radius:0 6px 6px 0;padding:10px 12px;margin:10px 0;'>"
                        f"<b style='font-size:0.85rem;color:{MTL_NAVY};'>{a.get('category','')}</b><br>"
                        f"<span style='font-size:0.82rem;color:#4a5568;'>{a.get('summary','')}</span></div>",
                        unsafe_allow_html=True)
                    if a.get("legal_issues"):
                        with st.expander("⚖ Vấn đề pháp lý", expanded=True):
                            for iss in a["legal_issues"]: st.markdown(f"- {iss}")
                    if a.get("relevant_laws"):
                        with st.expander("📋 Căn cứ pháp lý"):
                            for law in a["relevant_laws"]: st.markdown(f"`{law}`")
                    if a.get("action_items"):
                        with st.expander("✅ Hành động", expanded=True):
                            for i, act in enumerate(a["action_items"],1): st.markdown(f"{i}. {act}")
                    risk = a.get("risk_level","")
                    risk_ic = {"Cao":"🔴","Trung bình":"🟡","Thấp":"🟢"}.get(risk,"")
                    st.divider()
                    st.caption(f"Rủi ro: {risk_ic} {risk}")
                    bao_cao = (
                        f"VỤ VIỆC: {em['subject']}\nKHÁCH HÀNG: {em['fromName']}\n\n"
                        f"TÓM TẮT:\n{a.get('summary','')}\n\n"
                        f"VẤN ĐỀ PHÁP LÝ:\n" + "\n".join(f"- {i}" for i in a.get("legal_issues",[])) +
                        f"\n\nCĂN CỨ PHÁP LÝ:\n" + "\n".join(f"- {l}" for l in a.get("relevant_laws",[])) +
                        f"\n\nHÀNH ĐỘNG CẦN LÀM:\n" + "\n".join(f"{i+1}. {x}" for i,x in enumerate(a.get("action_items",[]))))
                    wb = tao_file_word("BÁO CÁO PHÂN TÍCH EMAIL", bao_cao, nd["ho_ten"], nd["chuc_vu"])
                    st.download_button("⬇️ Xuất báo cáo Word", data=wb,
                        file_name=f"PhanTichEmail_{datetime.now().strftime('%d%m%Y_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True)

            with ai2:
                tone_vi = {"formal":"Trang trọng","friendly":"Thân thiện",
                           "firm":"Kiên quyết","urgent":"Khẩn cấp"}
                tone_sel = st.radio("Giọng văn", options=list(tone_vi.keys()),
                                    format_func=lambda x: tone_vi[x],
                                    horizontal=True, key="ei_tone_r")
                st.session_state.ei_tone = tone_sel
                if st.button("✦ Tạo nháp AI", use_container_width=True, key="ei_gen"):
                    with st.spinner("Claude đang soạn..."):
                        st.session_state.ei_draft = soan_phan_hoi(
                            em, st.session_state.ei_analysis, tone_sel)
                    st.rerun()
                reply_to = st.text_input("Gửi đến", value=em.get("fromEmail",""), key="ei_to")
                draft = st.text_area("Nội dung phản hồi", value=st.session_state.ei_draft,
                                     height=240, key="ei_ta",
                                     placeholder="Nhấn '✦ Tạo nháp AI' hoặc tự soạn...")
                st.session_state.ei_draft = draft
                sa, sb = st.columns(2)
                with sa:
                    if draft.strip():
                        wb2 = tao_file_word(f"Phản hồi: {em['subject']}", draft, nd["ho_ten"], nd["chuc_vu"])
                        st.download_button("⬇️ Tải Word", data=wb2,
                            file_name=f"PhanHoi_{datetime.now().strftime('%d%m%Y_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True)
                with sb:
                    if st.button("📤 Gửi Gmail", type="primary",
                                 use_container_width=True, key="ei_send"):
                        if not draft.strip():
                            st.warning("Nhập nội dung trước")
                        elif not _connected:
                            st.error("Đăng nhập Google ở thanh bên trước")
                        else:
                            with st.spinner("Đang gửi..."):
                                ok = gui_email_gmail_api(reply_to, em["subject"], draft)
                            if ok:
                                st.session_state.ei_sent.append({
                                    "to": reply_to, "subject": em["subject"],
                                    "body": draft, "time": datetime.now().strftime("%H:%M %d/%m"),
                                })
                                st.success("✅ Email đã gửi!")
                                st.session_state.ei_draft = ""
                                st.rerun()

            with ai3:
                sent = st.session_state.ei_sent
                if not sent:
                    st.info("Chưa có email nào được gửi")
                else:
                    for item in reversed(sent):
                        with st.expander(f"✅ {item['time']} → {item['to']}"):
                            st.markdown(f"**{item['subject']}**")
                            st.text(item["body"][:300] + ("…" if len(item["body"])>300 else ""))


# ══════════════════════════════════════════════════════════════
#  TAB 6 — QUẢN LÝ CÔNG VIỆC (Task Management)
#  • Danh sách task (CRUD) + đánh dấu hoàn thành
#  • Lịch công việc theo tuần
#  • Hiệu suất nhân viên: tuần / tháng / năm  (chỉ admin)
#  • Báo cáo tuần tự động — gửi Gmail Thứ 5 20:00
# ══════════════════════════════════════════════════════════════

# ── Session-state khởi tạo (prefix mtl_task_ tránh xung đột) ──
# ── 5 Task cứng bắt buộc mỗi tuần ──────────────────────────────
MANDATORY_TASKS = [
    {
        "idx": 0,
        "title": "Power of One: Gặp 1 cộng tác viên",
        "icon": "🤝",
        "desc": "Gặp gỡ, kết nối với ít nhất 1 cộng tác viên trong tuần.",
        "priority": "high",
    },
    {
        "idx": 1,
        "title": "Gặp 1 khách hàng cũ",
        "icon": "👤",
        "desc": "Thăm hỏi, chăm sóc hoặc trao đổi công việc với 1 khách hàng đã hợp tác.",
        "priority": "high",
    },
    {
        "idx": 2,
        "title": "Gặp 1 khách hàng mới",
        "icon": "🌟",
        "desc": "Tiếp cận, tư vấn hoặc gặp mặt 1 khách hàng tiềm năng mới.",
        "priority": "high",
    },
    {
        "idx": 3,
        "title": "Ký 01 hợp đồng / Thực hiện 1 dự án mới",
        "icon": "✍️",
        "desc": "Ký kết hợp đồng dịch vụ pháp lý hoặc khởi động 1 dự án mới trong tuần.",
        "priority": "high",
    },
    {
        "idx": 4,
        "title": "Học tập 01 giờ",
        "icon": "📚",
        "desc": "Dành ít nhất 1 giờ học tập: nghiên cứu luật, đọc tài liệu chuyên môn, tham gia webinar...",
        "priority": "medium",
    },
]

_task_defaults = {
    "mtl_tasks":          [],   # Danh sách task thường
    "mtl_task_edit_id":   None, # ID task đang sửa
    "mtl_rpt_text":       "",   # Nội dung báo cáo AI mới nhất
    "mtl_last_sent_week": "",   # Tuần đã gửi báo cáo (vd: "2026-W17")
    "mtl_perf_period":    "week",
    "mtl_mandatory_done": {},   # {"{week}_{user_id}_{idx}": True/False}
    "mtl_mandatory_notes":{},   # {"{week}_{user_id}_{idx}": "ghi chú"}
}
for _k, _v in _task_defaults.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

def mkey(week, user_id, idx):
    """Khóa lưu trạng thái task cứng."""
    return f"{week}__{user_id}__{idx}"

def is_mandatory_done(week, user_id, idx):
    return st.session_state.mtl_mandatory_done.get(mkey(week, user_id, idx), False)

def set_mandatory_done(week, user_id, idx, val):
    st.session_state.mtl_mandatory_done[mkey(week, user_id, idx)] = val

def get_mandatory_note(week, user_id, idx):
    return st.session_state.mtl_mandatory_notes.get(mkey(week, user_id, idx), "")

def set_mandatory_note(week, user_id, idx, note):
    st.session_state.mtl_mandatory_notes[mkey(week, user_id, idx)] = note

# ─── Danh sách thành viên lấy từ TAI_KHOAN ───────────────────
def lay_thanh_vien():
    return [
        {"id": tk, "ho_ten": info["ho_ten"], "chuc_vu": info["chuc_vu"]}
        for tk, info in TAI_KHOAN.items()
    ]

THANH_VIEN = lay_thanh_vien()

def ten_nv(user_id):
    for m in THANH_VIEN:
        if m["id"] == user_id:
            return m["ho_ten"]
    return "Chưa phân công"

# ─── Helpers task ─────────────────────────────────────────────
def task_gen_id():
    return f"T{datetime.now().strftime('%Y%m%d%H%M%S%f')[-10:]}"

def tasks_of_week(week_str):
    """Lọc task theo tuần ISO (vd: '2026-W17')."""
    try:
        yr, wn = week_str.split("-W")
        yr, wn = int(yr), int(wn)
        jan4   = datetime(yr, 1, 4)
        day1   = jan4 - timedelta(days=jan4.weekday())
        ws     = day1 + timedelta(weeks=wn - 1)
        we     = ws + timedelta(days=6, hours=23, minutes=59, seconds=59)
        result = []
        for t in st.session_state.mtl_tasks:
            try:
                td = datetime.strptime(t["date"], "%Y-%m-%d")
                if ws <= td <= we:
                    result.append(t)
            except Exception:
                pass
        return result
    except Exception:
        return st.session_state.mtl_tasks

def tasks_of_month(year, month):
    result = []
    for t in st.session_state.mtl_tasks:
        try:
            td = datetime.strptime(t["date"], "%Y-%m-%d")
            if td.year == year and td.month == month:
                result.append(t)
        except Exception:
            pass
    return result

def tasks_of_year(year):
    result = []
    for t in st.session_state.mtl_tasks:
        try:
            td = datetime.strptime(t["date"], "%Y-%m-%d")
            if td.year == year:
                result.append(t)
        except Exception:
            pass
    return result

def cur_week_str():
    n = datetime.now()
    return n.strftime("%G-W%V")

# ─── Tạo báo cáo tuần bằng Claude ────────────────────────────
def tao_bao_cao_tuan(week_str, week_tasks):
    tv_data = []
    for m in THANH_VIEN:
        mt   = [t for t in week_tasks if t.get("assignee") == m["id"]]
        done = sum(1 for t in mt if t.get("done"))
        # Task cứng
        mand_done = sum(1 for mt2 in MANDATORY_TASKS
                        if is_mandatory_done(week_str, m["id"], mt2["idx"]))
        mand_notes = []
        for mt2 in MANDATORY_TASKS:
            note = get_mandatory_note(week_str, m["id"], mt2["idx"])
            tick = "✓" if is_mandatory_done(week_str, m["id"], mt2["idx"]) else "○"
            mand_notes.append(f"  [{tick}] {mt2['title']}" + (f" → {note}" if note else ""))
        tv_data.append(
            f"• {m['ho_ten']}: {done}/{len(mt)} task thường | "
            f"Task bắt buộc: {mand_done}/{len(MANDATORY_TASKS)}\n"
            + "\n".join(mand_notes)
        )

    task_lines = "\n".join(
        f"- [{'✓' if t.get('done') else '○'}] {t['title']} "
        f"| {ten_nv(t.get('assignee',''))} "
        f"| {'Cao' if t.get('priority')=='high' else 'TB' if t.get('priority')=='medium' else 'Thấp'} "
        f"| {t.get('date','—')}"
        + (f" | Kết quả: {t['notes']}" if t.get('notes') and t.get('done') else "")
        for t in week_tasks
    ) or "Không có task nào trong tuần này."

    total  = len(week_tasks)
    done_n = sum(1 for t in week_tasks if t.get("done"))

    prompt = f"""Viết BÁO CÁO CÔNG VIỆC TUẦN {week_str} của {TEN_CONG_TY} gửi Ban Lãnh Đạo.

THỐNG KÊ TASK THƯỜNG: {total} task | {done_n} hoàn thành | {total - done_n} chưa xong

TIẾN ĐỘ 5 MỤC TIÊU BẮT BUỘC (Power of One) VÀ TASK THƯỜNG THEO TỪNG LUẬT SƯ:
{chr(10).join(tv_data) or 'Chưa có dữ liệu.'}

DANH SÁCH TASK THƯỜNG:
{task_lines}

Yêu cầu: Viết tiếng Việt, văn phong pháp lý chuyên nghiệp, gồm:
1. TỔNG KẾT THÀNH TÍCH TUẦN
2. TIẾN ĐỘ 5 MỤC TIÊU "POWER OF ONE" (nêu rõ ai hoàn thành, ai chưa)
3. ĐÁNH GIÁ HIỆU SUẤT TỪNG LUẬT SƯ
4. CÔNG VIỆC TỒN ĐỌNG & GIẢI PHÁP
5. KẾ HOẠCH CÔNG VIỆC TUẦN TỚI
Ngắn gọn, súc tích, phù hợp trình lãnh đạo."""

    return goi_claude(
        [{"role": "user", "content": prompt}],
        f"Bạn là trợ lý hành chính pháp lý tại {TEN_CONG_TY}. "
        f"Viết báo cáo chuyên nghiệp, có cấu trúc rõ ràng.",
    )

# ─── Gửi báo cáo qua Gmail ────────────────────────────────────
def gui_bao_cao_gmail(to_list, cc_list, subject, body):
    """Gửi báo cáo đến nhiều người nhận TO + CC."""
    svc = _gmail_service()
    if not svc:
        return False, "Chưa kết nối Gmail"
    try:
        from email.mime.text import MIMEText
        msg = MIMEText(body, "plain", "utf-8")
        sender = st.session_state.get(f"gemail_{nd['ten_tk']}", "me")
        msg["From"]    = sender
        msg["To"]      = ", ".join(to_list)
        if cc_list:
            msg["Cc"]  = ", ".join(cc_list)
        msg["Subject"] = subject
        raw = base64.urlsafe_b64encode(msg.as_bytes()).decode()
        svc.users().messages().send(userId="me", body={"raw": raw}).execute()
        return True, "OK"
    except Exception as e:
        return False, str(e)

# ─── Biểu đồ cột HTML ─────────────────────────────────────────
def ve_bieu_do(labels, done_vals, todo_vals, title=""):
    """Render biểu đồ cột chồng (done=xanh, todo=vàng) bằng HTML thuần."""
    max_v = max(max(done_vals + todo_vals, default=1), 1)
    bar_h = 160  # px chiều cao tối đa

    bars_html = ""
    for i, lbl in enumerate(labels):
        d = done_vals[i] if i < len(done_vals) else 0
        u = todo_vals[i]  if i < len(todo_vals) else 0
        ph_d = int((d / max_v) * bar_h)
        ph_u = int((u / max_v) * bar_h)
        val_lbl = f"{d}/{d+u}" if (d + u) > 0 else ""
        bars_html += f"""
<div style="display:flex;flex-direction:column;align-items:center;flex:1;min-width:0;gap:2px;cursor:default;"
     title="{lbl}: {d} hoàn thành / {d+u} tổng">
  <div style="font-size:10px;font-weight:600;color:#555;">{val_lbl}</div>
  <div style="display:flex;flex-direction:column;justify-content:flex-end;height:{bar_h}px;width:100%;gap:1px;">
    <div style="background:#0f6e56;border-radius:3px 3px 0 0;height:{ph_d}px;width:100%;"></div>
    <div style="background:#C9A96E;border-radius:3px 3px 0 0;height:{ph_u}px;width:100%;"></div>
  </div>
  <div style="font-size:10px;color:#666;text-align:center;white-space:nowrap;
              overflow:hidden;text-overflow:ellipsis;width:100%;">{lbl}</div>
</div>"""

    legend = (
        "<span style='display:inline-flex;align-items:center;gap:4px;font-size:11px;color:#555;margin-right:12px;'>"
        "<span style='width:10px;height:10px;border-radius:2px;background:#0f6e56;display:inline-block;'></span>Hoàn thành</span>"
        "<span style='display:inline-flex;align-items:center;gap:4px;font-size:11px;color:#555;'>"
        "<span style='width:10px;height:10px;border-radius:2px;background:#C9A96E;display:inline-block;'></span>Chưa xong</span>"
    )

    return f"""
<div style="margin-bottom:6px;font-size:12px;font-weight:600;color:{MTL_NAVY};">{title}</div>
<div style="margin-bottom:8px;">{legend}</div>
<div style="display:flex;gap:4px;align-items:flex-end;height:{bar_h+40}px;padding:0 4px;">
  {bars_html}
</div>"""


# ══════════════════════════════════════════════════════════════
#  RENDER TAB 6
# ══════════════════════════════════════════════════════════════
with tab6:
    is_admin = (nd.get("vai_tro") == "quan_tri")

    # ── Banner ──
    st.markdown(f"""
<div style="background:linear-gradient(135deg,{MTL_NAVY2} 0%,{MTL_NAVY} 100%);
border-radius:10px;padding:14px 20px;margin-bottom:18px;
border-left:4px solid {MTL_GOLD};display:flex;align-items:center;justify-content:space-between;">
  <div>
    <span style="color:white;font-size:1.05rem;font-weight:700;">📌 Quản lý Công việc</span>
    <span style="color:{MTL_GOLD2};font-size:0.8rem;margin-left:12px;">
      Task · Lịch · Hiệu suất · Báo cáo tuần tự động
    </span>
  </div>
  <div style="background:rgba(168,135,74,0.2);border:1px solid {MTL_GOLD}55;
  border-radius:8px;padding:5px 12px;font-size:0.78rem;color:{MTL_GOLD2};">
    📅 Gửi tự động: Thứ 5 · 20:00
  </div>
</div>""", unsafe_allow_html=True)

    # ── Sub-tabs ──
    if is_admin:
        stab_labels = ["📋 Task", "📅 Lịch tuần", "📊 Hiệu suất", "📤 Báo cáo & Gửi"]
    else:
        stab_labels = ["📋 Task của tôi", "📅 Lịch tuần", "📤 Báo cáo & Gửi"]

    stabs = st.tabs(stab_labels)

    # ════════════════════════════════════════════
    #  STAB 0 — DANH SÁCH TASK
    # ════════════════════════════════════════════
    with stabs[0]:
        cur_wk_now = cur_week_str()

        # ── PHẦN 1: 5 TASK CỨNG BẮT BUỘC ───────────────────────
        st.markdown(f"""
<div style="background:linear-gradient(135deg,{MTL_NAVY2} 0%,{MTL_NAVY} 100%);
border-radius:10px;padding:12px 18px;margin-bottom:14px;
border-left:4px solid {MTL_GOLD};display:flex;align-items:center;gap:12px;">
  <span style="font-size:1.1rem;">🎯</span>
  <div>
    <span style="color:white;font-weight:700;font-size:0.95rem;">Task bắt buộc tuần — Power of One</span>
    <span style="color:{MTL_GOLD2};font-size:0.78rem;margin-left:10px;">5 mục tiêu cốt lõi · Tuần {cur_wk_now}</span>
  </div>
</div>""", unsafe_allow_html=True)

        # Xác định user nào đang xem (admin xem tất cả, ls xem mình)
        if is_admin:
            mandatory_users = THANH_VIEN
        else:
            mandatory_users = [m for m in THANH_VIEN if m["id"] == nd["ten_tk"]]

        for m_user in mandatory_users:
            uid = m_user["id"]
            done_count = sum(1 for mt in MANDATORY_TASKS
                             if is_mandatory_done(cur_wk_now, uid, mt["idx"]))

            if is_admin:
                pct = int(done_count / len(MANDATORY_TASKS) * 100)
                bar_c = "#0f6e56" if pct == 100 else (MTL_GOLD if pct >= 60 else "#e24b4a")
                st.markdown(
                    f"<div style='display:flex;align-items:center;gap:10px;margin:6px 0 4px;'>"
                    f"<div style='width:26px;height:26px;border-radius:50%;background:{MTL_NAVY};"
                    f"color:{MTL_GOLD2};display:flex;align-items:center;justify-content:center;"
                    f"font-size:11px;font-weight:700;flex-shrink:0;'>{m_user['ho_ten'][0]}</div>"
                    f"<span style='font-size:0.83rem;font-weight:600;color:{MTL_NAVY};'>"
                    f"{m_user['ho_ten']}</span>"
                    f"<div style='flex:1;background:#e0e8f5;border-radius:4px;height:5px;overflow:hidden;'>"
                    f"<div style='background:{bar_c};width:{pct}%;height:5px;'></div></div>"
                    f"<span style='font-size:0.75rem;color:#666;min-width:36px;text-align:right;'>"
                    f"{done_count}/{len(MANDATORY_TASKS)}</span>"
                    f"</div>",
                    unsafe_allow_html=True,
                )

            for mt in MANDATORY_TASKS:
                k_done = f"mand_cb_{uid}_{mt['idx']}_{cur_wk_now}"
                k_note = f"mand_note_{uid}_{mt['idx']}_{cur_wk_now}"
                k_exp  = f"mand_exp_{uid}_{mt['idx']}_{cur_wk_now}"
                done   = is_mandatory_done(cur_wk_now, uid, mt["idx"])
                note   = get_mandatory_note(cur_wk_now, uid, mt["idx"])

                pri_color = MTL_NAVY if mt["priority"] == "high" else MTL_GOLD
                done_style = "opacity:0.55;" if done else ""
                strike     = "text-decoration:line-through;color:#999;" if done else f"color:{MTL_NAVY};"

                # Chỉ cho phép tick nếu đúng user (hoặc admin có thể tick mọi người)
                can_tick = is_admin or (uid == nd["ten_tk"])

                row1, row2, row3 = st.columns([0.5, 7.5, 2])
                with row1:
                    if can_tick:
                        if st.button(
                            "☑" if done else "☐",
                            key=k_done,
                            help="Đánh dấu hoàn thành",
                        ):
                            set_mandatory_done(cur_wk_now, uid, mt["idx"], not done)
                            st.rerun()
                    else:
                        st.markdown(
                            f"<div style='padding:4px;color:{'#0f6e56' if done else '#ccc'};font-size:1rem;'>{'☑' if done else '☐'}</div>",
                            unsafe_allow_html=True,
                        )

                with row2:
                    st.markdown(
                        f"<div style='{done_style}display:flex;align-items:center;gap:8px;padding:3px 0;'>"
                        f"<span style='font-size:1rem;'>{mt['icon']}</span>"
                        f"<div>"
                        f"<span style='{strike}font-size:0.88rem;font-weight:600;'>{mt['title']}</span>"
                        f"<span style='background:{pri_color}22;color:{pri_color};font-size:10px;"
                        f"border-radius:4px;padding:1px 6px;margin-left:6px;font-weight:600;'>BẮT BUỘC</span>"
                        + (f"<span style='background:#eaf3de;color:#27500a;font-size:10px;"
                           f"border-radius:4px;padding:1px 6px;margin-left:4px;'>✓ Xong</span>" if done else "")
                        + f"<div style='font-size:0.75rem;color:#888;margin-top:1px;'>{mt['desc']}</div>"
                        + (f"<div style='font-size:0.76rem;color:#555;font-style:italic;margin-top:2px;'>"
                           f"📝 {note}</div>" if note else "")
                        + f"</div></div>",
                        unsafe_allow_html=True,
                    )

                with row3:
                    if can_tick:
                        with st.expander("📝 Ghi chú", expanded=False):
                            new_note = st.text_input(
                                "Kết quả",
                                value=note,
                                key=k_note,
                                placeholder="Ghi chú kết quả thực hiện...",
                                label_visibility="collapsed",
                            )
                            if new_note != note:
                                set_mandatory_note(cur_wk_now, uid, mt["idx"], new_note)
                                st.rerun()

                st.markdown(
                    "<div style='height:1px;background:#f0f3fa;margin:2px 0;'></div>",
                    unsafe_allow_html=True,
                )

            if is_admin and len(mandatory_users) > 1:
                st.markdown("<br>", unsafe_allow_html=True)

        # ── Tổng tiến độ 5 task bắt buộc của mình ──
        my_done_count = sum(1 for mt in MANDATORY_TASKS
                            if is_mandatory_done(cur_wk_now, nd["ten_tk"], mt["idx"]))
        pct_my = int(my_done_count / len(MANDATORY_TASKS) * 100)
        bar_color_my = "#0f6e56" if pct_my == 100 else (MTL_GOLD if pct_my >= 60 else "#e24b4a")
        status_lbl = "🎉 Hoàn thành xuất sắc!" if pct_my == 100 else f"Còn {len(MANDATORY_TASKS)-my_done_count} mục chưa xong"

        st.markdown(
            f"<div style='background:#f8f9fc;border:1px solid #e0e8f5;border-radius:8px;"
            f"padding:10px 14px;margin:8px 0 16px;display:flex;align-items:center;gap:14px;'>"
            f"<div style='flex:1;'>"
            f"<div style='font-size:0.8rem;color:#555;margin-bottom:5px;font-weight:600;'>"
            f"Tiến độ task bắt buộc tuần của bạn: <span style='color:{bar_color_my};'>{status_lbl}</span></div>"
            f"<div style='background:#e0e8f5;border-radius:4px;height:8px;overflow:hidden;'>"
            f"<div style='background:{bar_color_my};width:{pct_my}%;height:8px;"
            f"border-radius:4px;transition:width .3s;'></div></div></div>"
            f"<div style='font-size:1.4rem;font-weight:700;color:{bar_color_my};min-width:40px;text-align:right;'>"
            f"{pct_my}%</div></div>",
            unsafe_allow_html=True,
        )

        st.markdown(
            f"<div style='font-weight:700;font-size:0.95rem;color:{MTL_NAVY};"
            f"margin-bottom:10px;border-left:3px solid {MTL_GOLD};padding-left:10px;'>"
            f"📋 Task công việc thường</div>",
            unsafe_allow_html=True,
        )

        # Lọc task theo role
        all_tasks = st.session_state.mtl_tasks
        if is_admin:
            view_tasks = all_tasks
        else:
            view_tasks = [t for t in all_tasks if t.get("assignee") == nd["ten_tk"]]

        # ── Thanh thống kê ──
        total_t = len(view_tasks)
        done_t  = sum(1 for t in view_tasks if t.get("done"))
        todo_t  = total_t - done_t
        hi_t    = sum(1 for t in view_tasks if t.get("priority") == "high" and not t.get("done"))

        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Tổng task", total_t)
        c2.metric("Đã hoàn thành", done_t, delta=f"{int(done_t/total_t*100) if total_t else 0}%")
        c3.metric("Chưa xong", todo_t)
        c4.metric("🔴 Ưu tiên cao", hi_t)

        st.markdown("---")

        # ── Bộ lọc ──
        fc1, fc2, fc3, fc4 = st.columns([2, 1.5, 1.5, 1])
        with fc1:
            search_q = st.text_input("🔍 Tìm task", placeholder="Tên task...", label_visibility="collapsed", key="task_search")
        with fc2:
            if is_admin:
                nv_opts = ["Tất cả"] + [m["ho_ten"] for m in THANH_VIEN]
                filter_nv = st.selectbox("Nhân viên", nv_opts, label_visibility="collapsed", key="task_fnv")
            else:
                filter_nv = "Tất cả"
        with fc3:
            filter_status = st.selectbox("Trạng thái", ["Tất cả", "Chưa xong", "Đã xong"],
                                         label_visibility="collapsed", key="task_fst")
        with fc4:
            if st.button("➕ Thêm task", type="primary", use_container_width=True, key="task_add_btn"):
                st.session_state.mtl_task_edit_id = "__new__"
                st.rerun()

        # ── Form thêm / sửa task ──
        edit_id = st.session_state.mtl_task_edit_id
        if edit_id:
            editing_task = None
            if edit_id != "__new__":
                editing_task = next((t for t in all_tasks if t["id"] == edit_id), None)

            with st.expander(
                "✏️ Sửa task" if editing_task else "➕ Thêm task mới",
                expanded=True,
            ):
                with st.form("task_form", clear_on_submit=False):
                    ef1, ef2 = st.columns(2)
                    with ef1:
                        tf_title = st.text_input("Tiêu đề *",
                                                  value=editing_task["title"] if editing_task else "",
                                                  placeholder="Nhập tiêu đề task...")
                        tf_date  = st.date_input("Ngày thực hiện",
                                                  value=datetime.strptime(editing_task["date"], "%Y-%m-%d").date()
                                                  if editing_task and editing_task.get("date") else datetime.now().date())
                        tf_priority = st.selectbox("Độ ưu tiên",
                                                    ["high", "medium", "low"],
                                                    format_func=lambda x: {"high":"🔴 Cao","medium":"🟡 Trung bình","low":"🟢 Thấp"}[x],
                                                    index={"high":0,"medium":1,"low":2}.get(
                                                        editing_task.get("priority","medium") if editing_task else "medium", 1))
                    with ef2:
                        nv_ids   = [m["id"] for m in THANH_VIEN]
                        nv_names = [m["ho_ten"] for m in THANH_VIEN]
                        cur_nv   = editing_task.get("assignee", nd["ten_tk"]) if editing_task else nd["ten_tk"]
                        nv_idx   = nv_ids.index(cur_nv) if cur_nv in nv_ids else 0
                        tf_assignee = st.selectbox("Người thực hiện", nv_ids,
                                                    format_func=lambda x: ten_nv(x),
                                                    index=nv_idx)
                        tf_status   = st.selectbox("Trạng thái",
                                                    ["todo", "done"],
                                                    format_func=lambda x: "✅ Đã hoàn thành" if x=="done" else "⏳ Chưa xong",
                                                    index=1 if (editing_task and editing_task.get("done")) else 0)
                        tf_add_cal  = st.checkbox("📆 Thêm vào Google Calendar",
                                                   value=True if not editing_task else False)
                    tf_desc  = st.text_area("Mô tả", value=editing_task.get("desc","") if editing_task else "",
                                             placeholder="Mô tả chi tiết công việc...", height=70)
                    tf_notes = st.text_area("Ghi chú / Kết quả",
                                             value=editing_task.get("notes","") if editing_task else "",
                                             placeholder="Ghi chú sau khi hoàn thành...", height=60)

                    sb1, sb2, sb3 = st.columns([1,1,3])
                    with sb1:
                        saved = st.form_submit_button("💾 Lưu", type="primary", use_container_width=True)
                    with sb2:
                        cancelled = st.form_submit_button("Hủy", use_container_width=True)

                    if cancelled:
                        st.session_state.mtl_task_edit_id = None
                        st.rerun()

                    if saved:
                        if not tf_title.strip():
                            st.error("Vui lòng nhập tiêu đề task!")
                        else:
                            now_str = datetime.now().isoformat()
                            is_done = (tf_status == "done")
                            if editing_task:
                                for t in st.session_state.mtl_tasks:
                                    if t["id"] == edit_id:
                                        t.update({
                                            "title":      tf_title.strip(),
                                            "desc":       tf_desc.strip(),
                                            "assignee":   tf_assignee,
                                            "priority":   tf_priority,
                                            "date":       tf_date.strftime("%Y-%m-%d"),
                                            "notes":      tf_notes.strip(),
                                            "done":       is_done,
                                            "updated_at": now_str,
                                        })
                                        if is_done and not t.get("completed_at"):
                                            t["completed_at"] = now_str
                                        break
                            else:
                                new_task = {
                                    "id":         task_gen_id(),
                                    "title":      tf_title.strip(),
                                    "desc":       tf_desc.strip(),
                                    "assignee":   tf_assignee,
                                    "priority":   tf_priority,
                                    "date":       tf_date.strftime("%Y-%m-%d"),
                                    "notes":      tf_notes.strip(),
                                    "done":       is_done,
                                    "created_at": now_str,
                                    "updated_at": now_str,
                                    "completed_at": now_str if is_done else None,
                                    "cal_added":  False,
                                }
                                st.session_state.mtl_tasks.insert(0, new_task)

                                # Thêm Google Calendar nếu được chọn
                                if tf_add_cal and _gmail_service():
                                    try:
                                        from googleapiclient.discovery import build as _build
                                        svc_cal = _build("calendar", "v3",
                                                         credentials=st.session_state.get(f"gcred_{nd['ten_tk']}"))
                                        event = {
                                            "summary":     tf_title.strip(),
                                            "description": tf_desc.strip(),
                                            "start": {"date": tf_date.strftime("%Y-%m-%d")},
                                            "end":   {"date": tf_date.strftime("%Y-%m-%d")},
                                            "attendees": [{"email": st.session_state.get(
                                                f"gemail_{nd['ten_tk']}", "")}],
                                        }
                                        svc_cal.events().insert(
                                            calendarId="primary", body=event
                                        ).execute()
                                        new_task["cal_added"] = True
                                    except Exception:
                                        pass  # Calendar không bắt buộc

                            st.session_state.mtl_task_edit_id = None
                            st.success("✅ Đã lưu task!")
                            st.rerun()

        # ── Lọc danh sách ──
        filtered = view_tasks
        if search_q.strip():
            filtered = [t for t in filtered if search_q.lower() in t["title"].lower()
                        or search_q.lower() in t.get("desc","").lower()]
        if filter_nv != "Tất cả":
            nv_id = next((m["id"] for m in THANH_VIEN if m["ho_ten"] == filter_nv), None)
            if nv_id:
                filtered = [t for t in filtered if t.get("assignee") == nv_id]
        if filter_status == "Chưa xong":
            filtered = [t for t in filtered if not t.get("done")]
        elif filter_status == "Đã xong":
            filtered = [t for t in filtered if t.get("done")]

        st.caption(f"{len(filtered)} task")

        # ── Hiển thị từng task ──
        PRI_MAP  = {"high": ("🔴", "Cao"), "medium": ("🟡", "TB"), "low": ("🟢", "Thấp")}
        if not filtered:
            st.info("Chưa có task nào. Nhấn ➕ Thêm task để bắt đầu.")
        else:
            for task in filtered:
                icon_p, lbl_p = PRI_MAP.get(task.get("priority","medium"), ("🟡","TB"))
                done_icon = "✅" if task.get("done") else "⏳"
                cal_badge = " 📆" if task.get("cal_added") else ""

                with st.container():
                    col_cb, col_info, col_act = st.columns([0.5, 7, 2.5])

                    with col_cb:
                        # Nút tick hoàn thành
                        if st.button(
                            "☑" if task.get("done") else "☐",
                            key=f"cb_{task['id']}",
                            help="Đánh dấu hoàn thành / chưa xong",
                        ):
                            for t in st.session_state.mtl_tasks:
                                if t["id"] == task["id"]:
                                    t["done"] = not t["done"]
                                    if t["done"]:
                                        t["completed_at"] = datetime.now().isoformat()
                                    else:
                                        t["completed_at"] = None
                                    break
                            st.rerun()

                    with col_info:
                        title_style = "text-decoration:line-through;color:#999;" if task.get("done") else ""
                        completed_lbl = ""
                        if task.get("done") and task.get("completed_at"):
                            try:
                                dt_c = datetime.fromisoformat(task["completed_at"])
                                completed_lbl = f" · Xong {dt_c.strftime('%d/%m')}"
                            except Exception:
                                pass

                        st.markdown(
                            f"<div style='{title_style}font-weight:600;font-size:0.92rem;'>"
                            f"{done_icon} {task['title']}</div>"
                            f"<div style='font-size:0.78rem;color:#666;margin-top:2px;'>"
                            f"👤 {ten_nv(task.get('assignee',''))} &nbsp;|&nbsp; "
                            f"{icon_p} {lbl_p} &nbsp;|&nbsp; 📅 {task.get('date','—')}"
                            f"{cal_badge}{completed_lbl}"
                            + (f" &nbsp;|&nbsp; <i>\"{task['notes'][:60]}{'…' if len(task.get('notes',''))>60 else ''}\"</i>"
                               if task.get("notes") and task.get("done") else "")
                            + "</div>",
                            unsafe_allow_html=True,
                        )

                    with col_act:
                        ba1, ba2 = st.columns(2)
                        with ba1:
                            if st.button("✏️", key=f"edit_{task['id']}", help="Sửa task"):
                                st.session_state.mtl_task_edit_id = task["id"]
                                st.rerun()
                        with ba2:
                            if st.button("🗑️", key=f"del_{task['id']}", help="Xóa task"):
                                st.session_state.mtl_tasks = [
                                    t for t in st.session_state.mtl_tasks if t["id"] != task["id"]
                                ]
                                st.rerun()

                st.markdown(
                    "<div style='height:1px;background:#f0f0f0;margin:2px 0 4px;'></div>",
                    unsafe_allow_html=True,
                )

    # ════════════════════════════════════════════
    #  STAB 1 — LỊCH CÔNG VIỆC
    # ════════════════════════════════════════════
    with stabs[1]:
        now = datetime.now()
        if "mtl_cal_offset" not in st.session_state:
            st.session_state.mtl_cal_offset = 0

        # Điều hướng tuần
        nav1, nav2, nav3, nav4 = st.columns([1, 2, 1, 4])
        with nav1:
            if st.button("‹ Tuần trước", use_container_width=True, key="cal_prev"):
                st.session_state.mtl_cal_offset -= 1
                st.rerun()
        with nav3:
            if st.button("Tuần sau ›", use_container_width=True, key="cal_next"):
                st.session_state.mtl_cal_offset += 1
                st.rerun()
        with nav2:
            if st.button("Hôm nay", use_container_width=True, key="cal_today"):
                st.session_state.mtl_cal_offset = 0
                st.rerun()

        offset     = st.session_state.mtl_cal_offset
        week_start = now - timedelta(days=now.weekday()) + timedelta(weeks=offset)
        week_start = week_start.replace(hour=0, minute=0, second=0, microsecond=0)

        st.markdown(
            f"<div style='text-align:center;font-weight:700;color:{MTL_NAVY};margin-bottom:12px;'>"
            f"📅 Tuần {week_start.strftime('%d/%m')} — "
            f"{(week_start + timedelta(days=6)).strftime('%d/%m/%Y')}</div>",
            unsafe_allow_html=True,
        )

        # Build calendar HTML
        day_names = ["Thứ 2", "Thứ 3", "Thứ 4", "Thứ 5", "Thứ 6", "Thứ 7", "Chủ nhật"]
        cal_html  = "<div style='display:grid;grid-template-columns:repeat(7,1fr);gap:6px;'>"

        for i in range(7):
            day      = week_start + timedelta(days=i)
            day_str  = day.strftime("%Y-%m-%d")
            is_today = (day.date() == now.date())
            day_tasks = [t for t in st.session_state.mtl_tasks if t.get("date") == day_str]

            # Nếu không phải admin, chỉ hiển thị task của mình
            if not is_admin:
                day_tasks = [t for t in day_tasks if t.get("assignee") == nd["ten_tk"]]

            hdr_bg  = MTL_NAVY if is_today else "#f8f9fc"
            hdr_clr = "white"  if is_today else MTL_NAVY

            task_html = ""
            for t in day_tasks:
                bg_t   = "#0f6e56" if t.get("done") else MTL_NAVY
                txt_t  = "#9fe1cb" if t.get("done") else MTL_GOLD2
                strike = "text-decoration:line-through;" if t.get("done") else ""
                t_title  = t["title"]
                t_assign = ten_nv(t.get("assignee", ""))
                task_html += (
                    "<div style='background:" + bg_t + ";color:" + txt_t + ";border-radius:4px;"
                    "padding:3px 5px;font-size:10px;margin-bottom:3px;" + strike +
                    "white-space:nowrap;overflow:hidden;text-overflow:ellipsis;'"
                    " title='" + t_title + " - " + t_assign + "'>"
                    + t_title + "</div>"
                )

            border_style = "2px solid " + MTL_NAVY if is_today else "1px solid #e0e8f5"
            cal_html += (
                "<div style='border:" + border_style + ";"
                f"border-radius:8px;overflow:hidden;min-height:120px;'>"
                f"<div style='background:{hdr_bg};color:{hdr_clr};font-size:11px;"
                f"font-weight:600;padding:5px 7px;text-align:center;'>"
                f"{day_names[i]}<br>{day.strftime('%d/%m')}</div>"
                f"<div style='padding:6px;background:white;min-height:90px;'>"
                f"{task_html if task_html else '<span style=\"color:#ccc;font-size:10px;\">Trống</span>'}"
                f"</div></div>"
            )

        cal_html += "</div>"
        st.markdown(cal_html, unsafe_allow_html=True)

        st.markdown(
            "<div style='margin-top:10px;font-size:11px;color:#999;'>"
            "🟦 Task chưa xong &nbsp;|&nbsp; 🟩 Đã hoàn thành</div>",
            unsafe_allow_html=True,
        )

    # ════════════════════════════════════════════
    #  STAB 2 — HIỆU SUẤT (chỉ admin)
    # ════════════════════════════════════════════
    if is_admin:
        with stabs[2]:
            st.markdown(
                f"<div style='font-weight:700;font-size:1rem;color:{MTL_NAVY};"
                f"margin-bottom:16px;'>📊 Phân tích hiệu suất nhân viên</div>",
                unsafe_allow_html=True,
            )

            # Bộ lọc kỳ
            hp1, hp2, hp3 = st.columns([2, 1, 3])
            with hp1:
                period_sel = st.radio(
                    "Xem theo",
                    ["Tuần", "Tháng", "Năm"],
                    horizontal=True,
                    key="perf_period_radio",
                )
            with hp2:
                year_sel = st.selectbox(
                    "Năm",
                    list(range(datetime.now().year, datetime.now().year - 5, -1)),
                    key="perf_year_sel",
                )

            all_t = st.session_state.mtl_tasks

            # ── Số liệu tổng ──
            if period_sel == "Tuần":
                period_tasks = tasks_of_year(year_sel)
                title_suffix = f"năm {year_sel}"
            elif period_sel == "Tháng":
                period_tasks = tasks_of_year(year_sel)
                title_suffix = f"năm {year_sel}"
            else:
                period_tasks = tasks_of_year(year_sel)
                title_suffix = f"năm {year_sel}"

            done_all = sum(1 for t in period_tasks if t.get("done"))
            rate_all = int(done_all / len(period_tasks) * 100) if period_tasks else 0

            m1, m2, m3, m4 = st.columns(4)
            m1.metric(f"Task {title_suffix}", len(period_tasks))
            m2.metric("Hoàn thành", done_all, delta=f"{rate_all}%")
            m3.metric("Chưa xong", len(period_tasks) - done_all)
            m4.metric("Tỷ lệ TB", f"{rate_all}%")

            st.markdown("---")

            # ── Biểu đồ theo kỳ ──
            if period_sel == "Tuần":
                labels, done_v, todo_v = [], [], []
                for w in range(7, -1, -1):
                    ref = datetime.now() - timedelta(weeks=w)
                    wstr = ref.strftime("%G-W%V")
                    wt = tasks_of_week(wstr)
                    done_v.append(sum(1 for t in wt if t.get("done")))
                    todo_v.append(sum(1 for t in wt if not t.get("done")))
                    labels.append(f"T{ref.strftime('%V')}")
                chart_title = "Task 8 tuần gần nhất"

            elif period_sel == "Tháng":
                labels, done_v, todo_v = [], [], []
                for m in range(1, 13):
                    mt = tasks_of_month(year_sel, m)
                    done_v.append(sum(1 for t in mt if t.get("done")))
                    todo_v.append(sum(1 for t in mt if not t.get("done")))
                    labels.append(f"T{m}")
                chart_title = f"Task 12 tháng năm {year_sel}"

            else:  # Năm
                labels, done_v, todo_v = [], [], []
                for y in range(year_sel - 4, year_sel + 1):
                    yt = tasks_of_year(y)
                    done_v.append(sum(1 for t in yt if t.get("done")))
                    todo_v.append(sum(1 for t in yt if not t.get("done")))
                    labels.append(str(y))
                chart_title = "Task 5 năm gần nhất"

            st.markdown(
                f"<div style='background:white;border:1px solid #e0e8f5;border-radius:10px;"
                f"padding:16px 20px;margin-bottom:20px;'>"
                + ve_bieu_do(labels, done_v, todo_v, chart_title)
                + "</div>",
                unsafe_allow_html=True,
            )

            # ── Thẻ hiệu suất từng luật sư ──
            st.markdown(
                f"<div style='font-weight:600;color:{MTL_NAVY};margin-bottom:10px;'>"
                f"👤 Hiệu suất từng luật sư — {title_suffix}</div>",
                unsafe_allow_html=True,
            )

            member_rows = []
            for m in THANH_VIEN:
                mt   = [t for t in period_tasks if t.get("assignee") == m["id"]]
                done = sum(1 for t in mt if t.get("done"))
                rate = int(done / len(mt) * 100) if mt else 0
                hi   = sum(1 for t in mt if t.get("priority") == "high")
                member_rows.append((m, mt, done, rate, hi))

            member_rows.sort(key=lambda x: x[3], reverse=True)  # Sắp xếp theo tỷ lệ

            cols_nv = st.columns(2)
            for idx, (m, mt, done, rate, hi) in enumerate(member_rows):
                bar_color = "#0f6e56" if rate >= 80 else ("#C9A96E" if rate >= 50 else "#e24b4a")
                rank_icon = ["🥇","🥈","🥉"][idx] if idx < 3 else f"#{idx+1}"

                with cols_nv[idx % 2]:
                    st.markdown(
                        f"<div style='background:#f8f9fc;border:1px solid #e0e8f5;"
                        f"border-radius:10px;padding:14px;margin-bottom:12px;'>"
                        f"<div style='display:flex;align-items:center;gap:10px;margin-bottom:10px;'>"
                        f"<div style='width:34px;height:34px;border-radius:50%;background:{MTL_NAVY};"
                        f"color:{MTL_GOLD2};display:flex;align-items:center;justify-content:center;"
                        f"font-weight:700;font-size:14px;'>{m['ho_ten'][0]}</div>"
                        f"<div><div style='font-weight:600;font-size:0.88rem;'>{m['ho_ten']}</div>"
                        f"<div style='font-size:0.75rem;color:#888;'>{m['chuc_vu']}</div></div>"
                        f"<div style='margin-left:auto;font-size:1.1rem;'>{rank_icon}</div>"
                        f"</div>"
                        # Metrics
                        f"<div style='display:grid;grid-template-columns:repeat(3,1fr);"
                        f"gap:6px;margin-bottom:10px;text-align:center;'>"
                        f"<div style='background:white;border-radius:6px;padding:6px;"
                        f"border:1px solid #e0e8f5;'>"
                        f"<div style='font-size:1.2rem;font-weight:700;color:{MTL_NAVY};'>{len(mt)}</div>"
                        f"<div style='font-size:10px;color:#888;'>Tổng</div></div>"
                        f"<div style='background:white;border-radius:6px;padding:6px;"
                        f"border:1px solid #e0e8f5;'>"
                        f"<div style='font-size:1.2rem;font-weight:700;color:#0f6e56;'>{done}</div>"
                        f"<div style='font-size:10px;color:#888;'>Xong</div></div>"
                        f"<div style='background:white;border-radius:6px;padding:6px;"
                        f"border:1px solid #e0e8f5;'>"
                        f"<div style='font-size:1.2rem;font-weight:700;color:{bar_color};'>{rate}%</div>"
                        f"<div style='font-size:10px;color:#888;'>Tỷ lệ</div></div>"
                        f"</div>"
                        # Thanh tiến độ
                        f"<div style='background:#e0e8f5;border-radius:4px;height:6px;overflow:hidden;'>"
                        f"<div style='background:{bar_color};width:{rate}%;height:6px;"
                        f"border-radius:4px;'></div></div>"
                        f"<div style='display:flex;justify-content:space-between;"
                        f"font-size:10px;color:#999;margin-top:3px;'>"
                        f"<span>Hoàn thành</span><span>{rate}%</span></div>"
                        + (f"<div style='margin-top:6px;font-size:10px;color:#888;'>"
                           f"Ưu tiên cao: {hi} task</div>" if hi else "")
                        + f"</div>",
                        unsafe_allow_html=True,
                    )

            # ── Bảng xếp hạng ──
            st.markdown(
                f"<div style='font-weight:600;color:{MTL_NAVY};margin:16px 0 10px;'>"
                f"🏆 Bảng xếp hạng — {title_suffix}</div>",
                unsafe_allow_html=True,
            )

            rank_medals = ["🥇","🥈","🥉"]
            max_done    = max((r[2] for r in member_rows), default=1)
            table_html  = (
                f"<table style='width:100%;border-collapse:collapse;font-size:0.82rem;'>"
                f"<thead><tr style='background:{MTL_NAVY};color:white;'>"
                f"<th style='padding:8px 10px;text-align:left;'>#</th>"
                f"<th style='padding:8px 10px;text-align:left;'>Luật sư</th>"
                f"<th style='padding:8px 10px;text-align:center;'>Hoàn thành</th>"
                f"<th style='padding:8px 10px;text-align:center;'>Tổng</th>"
                f"<th style='padding:8px 10px;text-align:center;'>Tỷ lệ</th>"
                f"<th style='padding:8px 10px;'>Hiệu suất</th>"
                f"</tr></thead><tbody>"
            )
            for i, (m, mt, done, rate, hi) in enumerate(member_rows):
                bg_row = "#f8f9fc" if i % 2 == 0 else "white"
                medal  = rank_medals[i] if i < 3 else f"#{i+1}"
                rate_c = "#0f6e56" if rate >= 80 else ("#C9A96E" if rate >= 50 else "#e24b4a")
                bar_w  = int(done / max_done * 100) if max_done else 0
                table_html += (
                    f"<tr style='background:{bg_row};'>"
                    f"<td style='padding:8px 10px;'>{medal}</td>"
                    f"<td style='padding:8px 10px;font-weight:600;'>{m['ho_ten']}</td>"
                    f"<td style='padding:8px 10px;text-align:center;color:#0f6e56;"
                    f"font-weight:700;'>{done}</td>"
                    f"<td style='padding:8px 10px;text-align:center;color:#888;'>{len(mt)}</td>"
                    f"<td style='padding:8px 10px;text-align:center;'>"
                    f"<span style='background:{rate_c}22;color:{rate_c};border-radius:4px;"
                    f"padding:2px 7px;font-weight:600;'>{rate}%</span></td>"
                    f"<td style='padding:8px 10px;'>"
                    f"<div style='background:#e0e8f5;border-radius:3px;height:5px;width:120px;'>"
                    f"<div style='background:{rate_c};width:{bar_w}%;height:5px;"
                    f"border-radius:3px;'></div></div></td>"
                    f"</tr>"
                )
            table_html += "</tbody></table>"
            st.markdown(
                f"<div style='border:1px solid #e0e8f5;border-radius:10px;"
                f"overflow:hidden;'>{table_html}</div>",
                unsafe_allow_html=True,
            )

    # ════════════════════════════════════════════
    #  STAB 3 (admin) / STAB 2 (luật sư) — BÁO CÁO
    # ════════════════════════════════════════════
    rpt_tab_idx = 3 if is_admin else 2
    with stabs[rpt_tab_idx]:
        st.markdown(
            f"<div style='font-weight:700;font-size:0.95rem;color:{MTL_NAVY};"
            f"margin-bottom:16px;'>📤 Báo cáo công việc tuần — Gửi Gmail tự động Thứ 5 · 20:00</div>",
            unsafe_allow_html=True,
        )

        # Kiểm tra tự động gửi (Thứ 5, 20:xx)
        _now = datetime.now()
        _is_thu5_20h = (_now.weekday() == 3 and _now.hour == 20)
        _cur_wk = cur_week_str()
        _auto_triggered = (
            _is_thu5_20h and
            st.session_state.mtl_last_sent_week != _cur_wk and
            bool(st.session_state.mtl_tasks)
        )

        if _auto_triggered:
            st.warning("🤖 Phát hiện Thứ 5 lúc 20:00 — đang tự động tạo và gửi báo cáo tuần!")

        # Chọn tuần
        rc1, rc2 = st.columns([2, 3])
        with rc1:
            week_input = st.text_input(
                "Tuần báo cáo (ISO, vd: 2026-W17)",
                value=_cur_wk,
                key="rpt_week_input",
            )

        # Lấy task tuần được chọn
        week_tasks_rpt = tasks_of_week(week_input)
        done_rpt  = sum(1 for t in week_tasks_rpt if t.get("done"))
        total_rpt = len(week_tasks_rpt)

        r1, r2, r3, r4 = st.columns(4)
        r1.metric("Task tuần", total_rpt)
        r2.metric("Hoàn thành", done_rpt)
        r3.metric("Chưa xong", total_rpt - done_rpt)
        r4.metric("Tỷ lệ", f"{int(done_rpt/total_rpt*100) if total_rpt else 0}%")

        st.markdown("---")

        # Cấu hình email (chỉ hiển thị đầy đủ cho admin)
        if is_admin:
            with st.expander("📧 Cấu hình danh sách gửi", expanded=False):
                ec1, ec2 = st.columns(2)
                with ec1:
                    to_emails = st.text_area(
                        "Email nhận chính (TO) — mỗi dòng 1 email",
                        placeholder="manager@luatminhtu.vn\nbanlanhDao@luatminhtu.vn",
                        key="rpt_to_emails",
                        height=90,
                    )
                with ec2:
                    cc_emails = st.text_area(
                        "CC — mỗi dòng 1 email",
                        placeholder="director@luatminhtu.vn",
                        key="rpt_cc_emails",
                        height=90,
                    )
                rpt_subject = st.text_input(
                    "Tiêu đề email",
                    value=f"[Báo cáo tuần] {TEN_CONG_TY} — {week_input}",
                    key="rpt_subject",
                )
        else:
            to_emails   = st.session_state.get(f"gemail_{nd['ten_tk']}", "")
            cc_emails   = ""
            rpt_subject = f"[Báo cáo tuần] Công việc của tôi — {week_input}"

        # Nút tạo báo cáo
        rpt_c1, rpt_c2, rpt_c3 = st.columns([1, 1, 3])
        with rpt_c1:
            gen_btn = st.button("🤖 Tạo báo cáo AI", type="primary",
                                use_container_width=True, key="rpt_gen")
        with rpt_c2:
            send_btn_disabled = not bool(st.session_state.mtl_rpt_text.strip())
            send_btn = st.button(
                "📤 Gửi Gmail",
                use_container_width=True,
                disabled=send_btn_disabled,
                key="rpt_send",
            )

        # Tạo báo cáo
        if gen_btn or _auto_triggered:
            if not week_tasks_rpt and not _auto_triggered:
                st.warning("Không có task nào trong tuần được chọn.")
            else:
                with st.spinner("🤖 Claude đang tổng hợp báo cáo..."):
                    # Nếu không phải admin, lọc task của mình
                    tasks_for_rpt = week_tasks_rpt if is_admin else [
                        t for t in week_tasks_rpt if t.get("assignee") == nd["ten_tk"]
                    ]
                    rpt_text = tao_bao_cao_tuan(week_input, tasks_for_rpt)
                    st.session_state.mtl_rpt_text = rpt_text
                st.success("✅ Đã tạo báo cáo!")
                if _auto_triggered:
                    # Tự động gửi luôn
                    to_list  = [e.strip() for e in to_emails.splitlines() if e.strip()] if is_admin else [to_emails]
                    cc_list  = [e.strip() for e in cc_emails.splitlines() if e.strip()] if is_admin else []
                    ok, err  = gui_bao_cao_gmail(to_list, cc_list, rpt_subject, rpt_text)
                    if ok:
                        st.session_state.mtl_last_sent_week = _cur_wk
                        st.success(f"✅ Đã tự động gửi đến {', '.join(to_list)}")
                    else:
                        st.error(f"❌ Gửi thất bại: {err}")
                st.rerun()

        # Gửi thủ công
        if send_btn:
            if not _gmail_service():
                st.error("❌ Chưa kết nối Gmail. Đăng nhập Google ở thanh bên trước.")
            else:
                to_list = [e.strip() for e in to_emails.splitlines() if e.strip()] if is_admin else [to_emails]
                cc_list = [e.strip() for e in cc_emails.splitlines() if e.strip()] if is_admin else []
                if not any(to_list):
                    st.error("Hãy nhập ít nhất 1 email nhận chính (TO).")
                else:
                    with st.spinner("📤 Đang gửi..."):
                        ok, err = gui_bao_cao_gmail(
                            to_list, cc_list, rpt_subject,
                            st.session_state.mtl_rpt_text,
                        )
                    if ok:
                        st.session_state.mtl_last_sent_week = week_input
                        st.success(
                            f"✅ Đã gửi đến: **{', '.join(to_list)}**"
                            + (f" (CC: {', '.join(cc_list)})" if cc_list else "")
                        )
                    else:
                        st.error(f"❌ Gửi thất bại: {err}")

        # Hiển thị nội dung báo cáo
        if st.session_state.mtl_rpt_text:
            st.markdown("#### 📊 Nội dung báo cáo")
            rpt_edit = st.text_area(
                "Chỉnh sửa trước khi gửi (tùy chọn)",
                value=st.session_state.mtl_rpt_text,
                height=380,
                key="rpt_edit_area",
            )
            st.session_state.mtl_rpt_text = rpt_edit

            # Export Word
            if rpt_edit.strip():
                wb_rpt = tao_file_word(
                    f"BÁO CÁO TUẦN {week_input}",
                    rpt_edit,
                    nd["ho_ten"],
                    nd["chuc_vu"],
                )
                st.download_button(
                    "⬇️ Xuất Word",
                    data=wb_rpt,
                    file_name=f"BaoCaoTuan_{week_input.replace('-','_')}_{datetime.now().strftime('%d%m%Y')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

        # Trạng thái gửi tuần hiện tại
        st.markdown("---")
        last_sent = st.session_state.mtl_last_sent_week
        if last_sent:
            st.markdown(
                f"<div style='background:#eaf3de;border:1px solid #c0dd97;border-radius:8px;"
                f"padding:8px 12px;font-size:0.82rem;color:#27500a;'>"
                f"✅ Tuần <strong>{last_sent}</strong> đã gửi báo cáo thành công.</div>",
                unsafe_allow_html=True,
            )
        next_thu = _now + timedelta(days=(3 - _now.weekday()) % 7)
        st.caption(
            f"⏱ Gửi tự động tiếp theo: Thứ 5 {next_thu.strftime('%d/%m/%Y')} lúc 20:00 | "
            f"Hiện tại: {_now.strftime('%A %d/%m %H:%M')}"
        )
